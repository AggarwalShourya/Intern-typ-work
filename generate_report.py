#!/usr/bin/env python3
"""
Generate NeMo Hybrid ASR Latency Benchmark Report as a Word document.
"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# DATA
# ─────────────────────────────────────────────

# Sequential data
SCENARIOS_COMMON = [
    "baseline", "lang_hi", "lang_hi_en", "lang_te", "lang_te_en",
    "blank_penalty_2", "word_boost_default", "word_boost_strong",
    "word_boost_full", "lang+word_boost", "kw_boost_default",
    "kw_boost_strong", "kw_boost+lang_te_en", "timestamps",
    "timestamps+lang", "timestamps+wb", "full_pipeline"
]

CTC_RUN1 = {
    "mean": [57.2,58.1,58.1,59.2,56.7,56.4,58.5,60.7,59.9,62.0,69.7,69.5,73.0,58.0,60.8,60.4,75.8],
    "std":  [2.2,0.7,0.5,0.5,0.4,0.7,1.2,0.9,1.2,1.1,1.1,0.5,0.9,0.8,3.2,1.0,2.5],
    "p50":  [56.5,58.2,58.1,59.2,56.8,56.2,58.6,60.7,60.4,62.1,69.2,69.6,72.8,57.8,58.8,60.7,75.4],
    "p95":  [64.0,59.4,58.8,60.2,57.5,58.0,60.3,62.5,61.6,63.2,72.9,70.4,75.1,59.6,66.7,61.6,80.2],
    "p99":  [64.0,59.4,58.8,60.2,57.5,58.0,60.3,62.5,61.6,63.2,72.9,70.4,75.1,59.6,66.7,61.6,80.2],
    "rtf":  [0.006]*17,
}

TDT_RUN1_SCENARIOS = SCENARIOS_COMMON + ["mj_lang_hi","mj_lang_hi_en","mj_lang_te","mj_lang_te_en"]
TDT_RUN1 = {
    "mean": [124.3,143.1,141.4,149.9,132.6,129.7,134.2,135.0,132.5,140.0,138.3,139.1,144.7,128.5,132.8,133.4,150.0,182.5,129.7,180.5,129.7],
    "std":  [1.4,4.7,2.1,2.1,2.5,1.7,2.7,2.1,1.2,4.0,1.5,1.5,2.1,2.5,1.5,1.8,1.3,3.2,1.3,2.2,1.9],
    "p50":  [123.6,141.4,140.8,150.2,131.6,129.6,133.6,134.8,132.6,138.8,138.2,139.2,144.5,127.7,132.6,133.2,150.0,182.1,129.8,180.7,130.1],
    "p95":  [127.7,157.2,146.8,153.4,139.8,132.7,141.1,140.1,134.4,151.6,141.0,141.6,149.5,135.7,136.2,137.9,152.0,192.3,132.4,184.6,132.8],
    "p99":  [127.7,157.2,146.8,153.4,139.8,132.7,141.1,140.1,134.4,151.6,141.0,141.6,149.5,135.7,136.2,137.9,152.0,192.3,132.4,184.6,132.8],
    "rtf":  [0.012,0.014,0.014,0.015,0.013,0.013,0.013,0.014,0.013,0.014,0.014,0.014,0.014,0.013,0.013,0.013,0.015,0.018,0.013,0.018,0.013],
}

CTC_RUN2 = {
    "mean": [56.2,59.0,60.5,59.8,59.8,57.7,60.7,62.3,63.8,65.4,73.0,73.8,75.9,58.2,58.8,61.8,80.9],
    "std":  [0.8,0.7,0.5,0.5,0.6,0.8,0.9,1.0,1.4,0.9,1.1,0.9,0.9,1.1,0.6,1.1,2.5],
    "p50":  [56.1,59.0,60.5,59.7,59.7,57.6,61.1,62.5,63.8,65.4,72.7,74.0,75.7,58.1,58.7,62.0,79.9],
    "p95":  [57.5,60.6,61.7,60.7,60.5,59.4,61.9,63.4,66.1,66.7,75.6,75.1,78.2,60.6,60.3,64.1,84.8],
    "p99":  [57.5,60.6,61.7,60.7,60.5,59.4,61.9,63.4,66.1,66.7,75.6,75.1,78.2,60.6,60.3,64.1,84.8],
    "rtf":  [0.006]*17,
}

RNNT_RUN2_SCENARIOS = SCENARIOS_COMMON + ["mj_lang_hi","mj_lang_hi_en","mj_lang_te","mj_lang_te_en"]
RNNT_RUN2 = {
    "mean": [179.4,196.4,198.9,247.1,186.9,187.8,200.1,186.1,187.1,192.9,197.8,198.5,203.6,183.7,187.3,187.3,209.9,202.0,184.2,253.7,182.8],
    "std":  [3.1,4.1,2.9,4.1,4.9,2.7,30.4,2.0,3.3,4.8,3.0,2.1,2.5,3.6,6.0,2.1,3.7,5.2,2.7,4.9,2.3],
    "p50":  [179.8,195.7,198.4,245.6,186.2,187.6,186.3,186.3,186.7,191.6,197.7,198.5,203.2,183.0,185.1,186.8,211.1,200.0,184.4,251.8,182.7],
    "p95":  [186.4,209.6,203.7,255.6,200.6,192.2,261.3,190.6,197.2,207.5,204.1,203.0,208.6,192.0,206.3,190.7,217.7,215.9,189.4,269.1,187.7],
    "p99":  [186.4,209.6,203.7,255.6,200.6,192.2,261.3,190.6,197.2,207.5,204.1,203.0,208.6,192.0,206.3,190.7,217.7,215.9,189.4,269.1,187.7],
    "rtf":  [0.018,0.020,0.020,0.025,0.019,0.019,0.020,0.019,0.019,0.019,0.020,0.020,0.020,0.018,0.019,0.019,0.021,0.020,0.018,0.025,0.018],
}

# Concurrent
CONC_LEVELS = [2, 4, 8]
CONC_DATA = {
    "CTC (Run 1)":   {"latency": [71.4, 123.0, 219.4], "throughput": [27.94, 32.39, 36.08]},
    "TDT (Run 1)":   {"latency": [136.7, 175.9, 344.5], "throughput": [14.62, 22.68, 23.10]},
    "CTC (Run 2)":   {"latency": [78.9, 122.6, 242.9], "throughput": [25.29, 32.47, 32.63]},
    "RNNT (Run 2)":  {"latency": [188.8, 221.2, 453.9], "throughput": [10.58, 18.05, 17.54]},
    "CTC (Run 3)":   {"latency": [51.0, 81.4, 143.1], "throughput": [39.06, 46.93, 50.90]},
    "TDT (Run 3)":   {"latency": [63.4, 126.6, 166.3], "throughput": [31.44, 29.70, 41.78]},
}

# Poisson
POISSON_SCENARIOS = ["baseline", "lang_hi_en", "word_boost", "kw_boost", "full_pipeline"]
POISSON_DATA = {
    "CTC (Run 1)":  [58.6, 67.3, 67.6, 77.3, 81.1],
    "TDT (Run 1)":  [147.7, 164.2, 150.4, 161.3, 187.6],
    "CTC (Run 2)":  [62.0, 64.8, 67.4, 81.7, 112.7],
    "RNNT (Run 2)": [220.3, 296.9, 216.0, 226.4, 276.8],
}

# Stress
STRESS_DATA = {
    "CTC (Run 1)":  {"wall": 5.0,  "rate": 40.0, "mean": 2031, "p95": 3085},
    "TDT (Run 1)":  {"wall": 9.02, "rate": 22.2, "mean": 3695, "p95": 5676},
    "CTC (Run 2)":  {"wall": 5.20, "rate": 38.5, "mean": 2115, "p95": 3164},
    "RNNT (Run 2)": {"wall": 11.37,"rate": 17.6, "mean": 4667, "p95": 7186},
    "CTC (Run 3)":  {"wall": 4.11, "rate": 48.6, "mean": 1532, "p95": 2127},
    "TDT (Run 3)":  {"wall": 4.11, "rate": 48.7, "mean": 1582, "p95": 2259},
}

# Multi-duration
DURATIONS = [2, 5, 10, 15, 20, 25, 30]
MULTI_SOLO = {
    "CTC (Run 1)":  [32.9, 37.0, 44.4, 53.4, 63.4, 107.5, 131.9],
    "TDT (Run 1)":  [50.3, 82.2, 123.4, 167.2, 214.6, 340.8, 358.9],
    "CTC (Run 2)":  [35.7, 38.9, 47.4, 55.1, 65.1, 102.8, 111.8],
    "RNNT (Run 2)": [58.8, 101.5, 180.7, 247.8, 326.9, 481.8, 541.4],
    "CTC (Run 3)":  [33.8, 38.9, 46.7, 55.1, 64.9, 96.3, 104.2],
    "TDT (Run 3)":  [38.5, 53.3, 71.5, 97.8, 114.7, 189.8, 211.8],
}
MULTI_CONC = {
    "CTC (Run 1)":  [18.0, 41.9, 68.4, 102.2, 160.3, 223.4, 315.1],
    "TDT (Run 1)":  [35.7, 107.8, 217.8, 368.8, 650.2, 1012.6, 1041.4],
    "CTC (Run 2)":  [21.0, 108.0, 196.1, 231.2, 415.4, 551.3, 595.6],
    "RNNT (Run 2)": [43.9, 248.5, 462.3, 764.3, 1157.9, 1630.4, 1903.6],
    "CTC (Run 3)":  [20.2, 41.7, 132.3, 292.9, 341.7, 403.1, 471.9],
    "TDT (Run 3)":  [23.3, 58.8, 167.9, 312.6, 558.6, 750.5, 747.3],
}

# masked_joint comparison
MJ_LABELS = ["hi (legacy)", "hi (mj_)", "hi_en (legacy)", "hi_en (mj_)","te (legacy)","te (mj_)","te_en (legacy)","te_en (mj_)"]
MJ_TDT    = [143.1, 182.5, 141.4, 129.7, 149.9, 180.5, 132.6, 129.7]
MJ_RNNT   = [196.4, 202.0, 198.9, 184.2, 247.1, 253.7, 186.9, 182.8]

# Run 3: CTC + TDT (Optimized)
CTC_RUN3 = {
    "mean": [60.5, 58.5, 56.9, 57.9, 57.4, 56.6, 60.9, 61.0, 61.5, 62.3, 64.8, 68.0, 67.3, 59.1, 57.6, 60.9, 66.9],
    "std":  [1.8, 1.3, 0.7, 0.8, 1.1, 0.4, 0.8, 0.5, 0.7, 0.5, 0.4, 3.3, 0.6, 1.0, 0.5, 0.8, 1.0],
    "p50":  [60.1, 58.8, 56.8, 58.0, 57.0, 56.6, 60.9, 60.9, 61.4, 62.2, 64.7, 66.8, 67.2, 58.5, 57.5, 60.6, 66.8],
    "p95":  [66.3, 60.6, 58.2, 60.1, 59.7, 57.3, 62.2, 62.1, 63.1, 63.1, 65.5, 76.4, 68.4, 61.0, 58.8, 62.4, 68.9],
    "p99":  [66.3, 60.6, 58.2, 60.1, 59.7, 57.3, 62.2, 62.1, 63.1, 63.1, 65.5, 76.4, 68.4, 61.0, 58.8, 62.4, 68.9],
    "rtf":  [0.006]*17,
}

TDT_RUN3_SCENARIOS = SCENARIOS_COMMON + ["mj_lang_hi", "mj_lang_hi_en", "mj_lang_te", "mj_lang_te_en"]
TDT_RUN3 = {
    "mean": [71.8, 75.5, 77.4, 105.7, 72.1, 71.2, 91.1, 90.2, 92.5, 88.2, 87.9, 88.5, 88.8, 80.8, 79.6, 91.1, 89.0, 92.2, 77.2, 96.2, 72.0],
    "std":  [2.0, 0.6, 0.5, 0.4, 0.5, 0.6, 1.4, 0.8, 0.5, 0.6, 0.6, 1.1, 0.6, 1.1, 0.4, 1.0, 0.4, 0.3, 0.5, 0.7, 0.4],
    "p50":  [71.4, 75.4, 77.5, 105.7, 72.1, 71.2, 90.8, 89.9, 92.3, 88.2, 87.8, 88.5, 88.8, 80.5, 79.7, 90.8, 88.9, 92.1, 77.2, 96.1, 71.9],
    "p95":  [78.7, 76.6, 78.9, 106.4, 73.5, 72.8, 93.7, 92.3, 93.6, 89.4, 89.4, 91.1, 90.0, 83.8, 80.2, 93.2, 89.8, 92.9, 78.2, 98.1, 73.0],
    "p99":  [78.7, 76.6, 78.9, 106.4, 73.5, 72.8, 93.7, 92.3, 93.6, 89.4, 89.4, 91.1, 90.0, 83.8, 80.2, 93.2, 89.8, 92.9, 78.2, 98.1, 73.0],
    "rtf":  [0.007, 0.008, 0.008, 0.011, 0.007, 0.007, 0.009, 0.009, 0.009, 0.009, 0.009, 0.009, 0.009, 0.008, 0.008, 0.009, 0.009, 0.009, 0.008, 0.010, 0.007],
}

MJ_TDT_RUN3 = [75.5, 92.2, 77.4, 77.2, 105.7, 96.2, 72.1, 72.0]

# Poisson Run 3 (20 req/s, not 4)
POISSON_RUN3_RATE = 20  # req/s (vs 4 req/s in Run 1/2)
POISSON_DATA_RUN3 = {
    "CTC (Run 3)": [58.3, 63.4, 71.4, 84.9, 111.0],
    "TDT (Run 3)": [103.6, 121.0, 208.3, 232.7, 206.1],
}

# ─────────────────────────────────────────────
# COLOR PALETTE
# ─────────────────────────────────────────────
COLORS = {
    "CTC (Run 1)":  "#2196F3",   # blue
    "TDT (Run 1)":  "#FF9800",   # orange
    "CTC (Run 2)":  "#4CAF50",   # green
    "RNNT (Run 2)": "#E91E63",   # pink/red
    "CTC (Run 3)":  "#00BCD4",   # teal
    "TDT (Run 3)":  "#9C27B0",   # purple
}
HEADING_COLOR = RGBColor(31, 73, 125)
ALT_ROW_COLOR = "DCE6F1"
HEADER_ROW_COLOR = "1F497D"

CHART_DIR = "/tmp/asr_charts"
os.makedirs(CHART_DIR, exist_ok=True)

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "axes.spines.top": False,
    "axes.spines.right": False,
})

try:
    plt.style.use("seaborn-v0_8-whitegrid")
except:
    try:
        plt.style.use("ggplot")
    except:
        pass

# ─────────────────────────────────────────────
# CHART GENERATION
# ─────────────────────────────────────────────

def save(fig, name):
    path = os.path.join(CHART_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


def chart1_sequential_bar():
    """Bar chart: sequential mean latency, 17 common scenarios, CTC vs TDT (R1) and CTC vs RNNT (R2)."""
    scenarios = SCENARIOS_COMMON
    x = np.arange(len(scenarios))
    w = 0.35

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(18, 6), sharey=False)

    # Run 1: CTC vs TDT
    bars1 = ax1.bar(x - w/2, CTC_RUN1["mean"], w, label="CTC", color=COLORS["CTC (Run 1)"], alpha=0.9)
    bars2 = ax1.bar(x + w/2, TDT_RUN1["mean"][:17], w, label="TDT", color=COLORS["TDT (Run 1)"], alpha=0.9)
    ax1.errorbar(x - w/2, CTC_RUN1["mean"], yerr=CTC_RUN1["std"], fmt="none", color="black", capsize=3, linewidth=1)
    ax1.errorbar(x + w/2, TDT_RUN1["mean"][:17], yerr=TDT_RUN1["std"][:17], fmt="none", color="black", capsize=3, linewidth=1)
    ax1.set_xticks(x)
    ax1.set_xticklabels(scenarios, rotation=45, ha="right", fontsize=8)
    ax1.set_ylabel("Mean Latency (ms)", fontsize=11)
    ax1.set_title("Run 1: CTC vs TDT Sequential Latency", fontsize=12, fontweight="bold")
    ax1.legend(fontsize=10)
    ax1.set_ylim(0, 200)
    ax1.grid(axis="y", alpha=0.4)

    # Run 2: CTC vs RNNT
    bars3 = ax2.bar(x - w/2, CTC_RUN2["mean"], w, label="CTC", color=COLORS["CTC (Run 2)"], alpha=0.9)
    bars4 = ax2.bar(x + w/2, RNNT_RUN2["mean"][:17], w, label="RNNT", color=COLORS["RNNT (Run 2)"], alpha=0.9)
    ax2.errorbar(x - w/2, CTC_RUN2["mean"], yerr=CTC_RUN2["std"], fmt="none", color="black", capsize=3, linewidth=1)
    ax2.errorbar(x + w/2, RNNT_RUN2["mean"][:17], yerr=RNNT_RUN2["std"][:17], fmt="none", color="black", capsize=3, linewidth=1)
    ax2.set_xticks(x)
    ax2.set_xticklabels(scenarios, rotation=45, ha="right", fontsize=8)
    ax2.set_ylabel("Mean Latency (ms)", fontsize=11)
    ax2.set_title("Run 2: CTC vs RNNT Sequential Latency", fontsize=12, fontweight="bold")
    ax2.legend(fontsize=10)
    ax2.set_ylim(0, 300)
    ax2.grid(axis="y", alpha=0.4)

    fig.suptitle("Sequential Latency — All 17 Common Scenarios (9.98s audio)", fontsize=14, fontweight="bold", y=1.01)
    fig.tight_layout()
    return save(fig, "chart1_sequential.png")


def chart2_concurrent():
    """Line chart: concurrent latency vs concurrency, dual y-axis."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(13, 5))

    for label, data in CONC_DATA.items():
        ax1.plot(CONC_LEVELS, data["latency"], marker="o", label=label,
                 color=COLORS[label], linewidth=2, markersize=7)
        ax2.plot(CONC_LEVELS, data["throughput"], marker="s", label=label,
                 color=COLORS[label], linewidth=2, markersize=7, linestyle="--")

    ax1.set_xlabel("Concurrency Level", fontsize=11)
    ax1.set_ylabel("Mean Latency (ms)", fontsize=11)
    ax1.set_title("Mean Latency vs Concurrency", fontsize=12, fontweight="bold")
    ax1.set_xticks(CONC_LEVELS)
    ax1.legend(fontsize=9)
    ax1.grid(alpha=0.4)

    ax2.set_xlabel("Concurrency Level", fontsize=11)
    ax2.set_ylabel("Throughput (req/s)", fontsize=11)
    ax2.set_title("Throughput vs Concurrency", fontsize=12, fontweight="bold")
    ax2.set_xticks(CONC_LEVELS)
    ax2.legend(fontsize=9)
    ax2.grid(alpha=0.4)

    fig.suptitle("Concurrent Benchmark — Latency & Throughput", fontsize=14, fontweight="bold")
    fig.tight_layout()
    return save(fig, "chart2_concurrent.png")


def chart3_poisson():
    """Grouped bar chart: Poisson mean latency for 5 scenarios."""
    scenarios = POISSON_SCENARIOS
    strategies = list(POISSON_DATA.keys())
    x = np.arange(len(scenarios))
    w = 0.18

    fig, ax = plt.subplots(figsize=(13, 6))
    for i, strat in enumerate(strategies):
        offset = (i - 1.5) * w
        ax.bar(x + offset, POISSON_DATA[strat], w, label=strat,
               color=list(COLORS.values())[i], alpha=0.88)

    ax.set_xticks(x)
    ax.set_xticklabels(scenarios, fontsize=11)
    ax.set_ylabel("Mean Latency (ms)", fontsize=12)
    ax.set_title("Poisson Traffic (4 req/s, 12s window) — Mean Latency by Scenario", fontsize=13, fontweight="bold")
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.4)
    fig.tight_layout()
    return save(fig, "chart3_poisson.png")


def chart4_stress():
    """Bar chart: stress test mean latency and throughput."""
    labels = list(STRESS_DATA.keys())
    means  = [STRESS_DATA[k]["mean"] for k in labels]
    p95s   = [STRESS_DATA[k]["p95"] for k in labels]
    rates  = [STRESS_DATA[k]["rate"] for k in labels]

    x = np.arange(len(labels))
    w = 0.3

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(13, 5))

    colors = [COLORS[k] for k in labels]
    b1 = ax1.bar(x - w/2, means, w, label="Mean Latency", color=colors, alpha=0.85)
    b2 = ax1.bar(x + w/2, p95s, w, label="P95 Latency", color=colors, alpha=0.5, hatch="//")
    ax1.set_xticks(x)
    ax1.set_xticklabels(labels, fontsize=10)
    ax1.set_ylabel("Latency (ms)", fontsize=11)
    ax1.set_title("Stress Test Latency\n(200 requests, 100 concurrent)", fontsize=12, fontweight="bold")
    mean_patch = mpatches.Patch(facecolor="gray", alpha=0.85, label="Mean")
    p95_patch  = mpatches.Patch(facecolor="gray", alpha=0.5, hatch="//", label="P95")
    ax1.legend(handles=[mean_patch, p95_patch], fontsize=10)
    ax1.grid(axis="y", alpha=0.4)

    ax2.bar(x, rates, color=colors, alpha=0.85)
    ax2.set_xticks(x)
    ax2.set_xticklabels(labels, fontsize=10)
    ax2.set_ylabel("Throughput (req/s)", fontsize=11)
    ax2.set_title("Stress Test Throughput\n(200 requests, 100 concurrent)", fontsize=12, fontweight="bold")
    ax2.grid(axis="y", alpha=0.4)
    for i, v in enumerate(rates):
        ax2.text(i, v + 0.5, f"{v}", ha="center", fontsize=10, fontweight="bold")

    fig.suptitle("Stress Test Results — All Strategies", fontsize=14, fontweight="bold")
    fig.tight_layout()
    return save(fig, "chart4_stress.png")


def chart5_multiduration():
    """Line chart: multi-duration solo latency vs audio duration."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

    for label, vals in MULTI_SOLO.items():
        ax1.plot(DURATIONS, vals, marker="o", label=label, color=COLORS[label], linewidth=2, markersize=7)
    ax1.axvline(20, color="red", linestyle="--", linewidth=1.5, label="Chunking threshold (20s)")
    ax1.set_xlabel("Audio Duration (s)", fontsize=11)
    ax1.set_ylabel("Mean Latency (ms)", fontsize=11)
    ax1.set_title("Solo Latency vs Audio Duration", fontsize=12, fontweight="bold")
    ax1.legend(fontsize=9)
    ax1.grid(alpha=0.4)
    ax1.set_xticks(DURATIONS)

    for label, vals in MULTI_CONC.items():
        ax2.plot(DURATIONS, vals, marker="s", label=label, color=COLORS[label], linewidth=2, markersize=7, linestyle="--")
    ax2.axvline(20, color="red", linestyle="--", linewidth=1.5, label="Chunking threshold (20s)")
    ax2.set_xlabel("Audio Duration (s)", fontsize=11)
    ax2.set_ylabel("Mean Latency (ms)", fontsize=11)
    ax2.set_title("Concurrent Latency vs Audio Duration", fontsize=12, fontweight="bold")
    ax2.legend(fontsize=9)
    ax2.grid(alpha=0.4)
    ax2.set_xticks(DURATIONS)

    fig.suptitle("Multi-Duration Batch — Solo vs Concurrent Latency", fontsize=14, fontweight="bold")
    fig.tight_layout()
    return save(fig, "chart5_multiduration.png")


def chart6_maskedjoint():
    """Bar chart: masked_joint vs legacy for TDT, RNNT, and TDT Run 3."""
    x = np.arange(len(MJ_LABELS))
    w = 0.25

    fig, ax = plt.subplots(figsize=(14, 6))
    ax.bar(x - w, MJ_TDT, w, label="TDT (Run 1)", color=COLORS["TDT (Run 1)"], alpha=0.88)
    ax.bar(x, MJ_RNNT, w, label="RNNT (Run 2)", color=COLORS["RNNT (Run 2)"], alpha=0.88)
    ax.bar(x + w, MJ_TDT_RUN3, w, label="TDT (Run 3, Optimized)", color=COLORS["TDT (Run 3)"], alpha=0.88)

    ax.set_xticks(x)
    ax.set_xticklabels(MJ_LABELS, rotation=20, ha="right", fontsize=10)
    ax.set_ylabel("Mean Latency (ms)", fontsize=12)
    ax.set_title("masked_joint (mj_) vs Legacy Language Constraint\nTDT, RNNT, and TDT Optimized", fontsize=13, fontweight="bold")
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.4)

    # Add value labels
    for i, v in enumerate(MJ_TDT):
        ax.text(i - w, v + 2, f"{v}", ha="center", fontsize=7)
    for i, v in enumerate(MJ_RNNT):
        ax.text(i, v + 2, f"{v}", ha="center", fontsize=7)
    for i, v in enumerate(MJ_TDT_RUN3):
        ax.text(i + w, v + 2, f"{v}", ha="center", fontsize=7)

    fig.tight_layout()
    return save(fig, "chart6_maskedjoint.png")


def chart7_optimization_comparison():
    """Grouped bar chart: TDT Run 1 vs TDT Run 3 for key scenarios."""
    labels = ["baseline", "full_pipeline", "mj_lang_hi", "mj_lang_te", "kw_boost_default", "word_boost_default"]
    # TDT Run 1 values for these scenarios
    tdt_r1_vals = [
        TDT_RUN1["mean"][0],   # baseline
        TDT_RUN1["mean"][16],  # full_pipeline
        TDT_RUN1["mean"][17],  # mj_lang_hi
        TDT_RUN1["mean"][19],  # mj_lang_te
        TDT_RUN1["mean"][10],  # kw_boost_default
        TDT_RUN1["mean"][6],   # word_boost_default
    ]
    # TDT Run 3 values for these scenarios
    tdt_r3_vals = [
        TDT_RUN3["mean"][0],   # baseline
        TDT_RUN3["mean"][16],  # full_pipeline
        TDT_RUN3["mean"][17],  # mj_lang_hi
        TDT_RUN3["mean"][19],  # mj_lang_te
        TDT_RUN3["mean"][10],  # kw_boost_default
        TDT_RUN3["mean"][6],   # word_boost_default
    ]

    x = np.arange(len(labels))
    w = 0.35

    fig, ax = plt.subplots(figsize=(12, 6))
    bars1 = ax.bar(x - w/2, tdt_r1_vals, w, label="TDT (Run 1)", color=COLORS["TDT (Run 1)"], alpha=0.88)
    bars2 = ax.bar(x + w/2, tdt_r3_vals, w, label="TDT (Run 3, Optimized)", color=COLORS["TDT (Run 3)"], alpha=0.88)

    # Add value labels
    for bar in bars1:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, h + 2, f"{h:.0f}", ha="center", fontsize=9, fontweight="bold")
    for bar in bars2:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, h + 2, f"{h:.0f}", ha="center", fontsize=9, fontweight="bold")

    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=20, ha="right", fontsize=10)
    ax.set_ylabel("Mean Latency (ms)", fontsize=12)
    ax.set_title("TDT Optimization Impact — Run 1 (Unoptimized) vs Run 3 (Optimized)", fontsize=13, fontweight="bold")
    ax.legend(fontsize=11)
    ax.grid(axis="y", alpha=0.4)
    ax.set_ylim(0, max(tdt_r1_vals) * 1.15)

    fig.tight_layout()
    return save(fig, "chart7_optimization.png")


# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_font(cell, bold=False, color=None, size=None, italic=False):
    for para in cell.paragraphs:
        for run in para.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color:
                run.font.color.rgb = color
            if size:
                run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.color.rgb = HEADING_COLOR
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    return para


def add_body(doc, text):
    para = doc.add_paragraph(text)
    para.style.font.size = Pt(10.5)
    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(text, style="List Bullet")
    para.style.font.size = Pt(10.5)
    return para


def add_sequential_table(doc, scenarios, data, label, alt_color=ALT_ROW_COLOR):
    headers = ["Scenario", "Mean (ms)", "Std", "P50", "P95", "P99", "RTF"]
    table = doc.add_table(rows=1 + len(scenarios), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        set_cell_bg(cell, HEADER_ROW_COLOR)
        set_cell_font(cell, bold=True, color=RGBColor(255, 255, 255), size=9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, sc in enumerate(scenarios):
        row = table.rows[r_idx + 1]
        vals = [
            sc,
            f"{data['mean'][r_idx]:.1f}",
            f"{data['std'][r_idx]:.1f}",
            f"{data['p50'][r_idx]:.1f}",
            f"{data['p95'][r_idx]:.1f}",
            f"{data['p99'][r_idx]:.1f}",
            f"{data['rtf'][r_idx]:.3f}",
        ]
        for c_idx, v in enumerate(vals):
            cell = row.cells[c_idx]
            cell.text = v
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(cell, alt_color)
            set_cell_font(cell, size=8.5)

    # Column widths
    widths = [Cm(3.5), Cm(2.0), Cm(1.5), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    return table


def set_page_size_a4(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def add_page_numbers(doc):
    """Add page numbers to footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_divider(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


# ─────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────

def build_doc(chart_paths):
    doc = Document()
    set_page_size_a4(doc)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    add_page_numbers(doc)

    # ── COVER PAGE ──────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("NeMo Hybrid ASR")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = HEADING_COLOR
    run.font.name = "Calibri"

    title_para2 = doc.add_paragraph()
    title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title_para2.add_run("Latency Benchmark Report")
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = HEADING_COLOR
    run2.font.name = "Calibri"

    doc.add_paragraph()
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("CTC vs TDT vs RNNT Decoding Strategies")
    sub_run.font.size = Pt(16)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(80, 80, 120)
    sub_run.font.name = "Calibri"

    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_rows = [
        ("Date", "April 2026"),
        ("Audio", "overall.wav — 9.98s, 16 kHz, multilingual (Telugu/English/Hindi)"),
        ("Model", "NeMo Hybrid ASR (CTC + TDT + RNNT heads)"),
        ("Server", "NVIDIA Triton Inference Server"),
        ("Test Runs", "Run 1: CTC + TDT   |   Run 2: CTC + RNNT   |   Run 3: CTC + TDT (Optimized)"),
    ]
    for i, (k, v) in enumerate(info_rows):
        row = info_table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_bg(row.cells[0], "1F497D")
        set_cell_font(row.cells[0], bold=True, color=RGBColor(255, 255, 255), size=10)
        if i % 2 == 0:
            set_cell_bg(row.cells[1], "DCE6F1")
        set_cell_font(row.cells[1], size=10)
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)

    doc.add_page_break()

    # ── SECTION 1: OVERVIEW ──────────────────────────
    add_heading(doc, "1. Overview", level=1)
    add_divider(doc)
    add_body(doc, (
        "This report presents a comprehensive latency benchmark of the NeMo Hybrid ASR server, "
        "covering three decoding strategies (CTC, TDT, RNNT) across five distinct test modes. "
        "The goal is to characterise per-request latency, throughput, scalability, and resilience "
        "under a variety of real-world conditions."
    ))
    doc.add_paragraph()

    add_heading(doc, "Audio Under Test", level=2)
    add_bullet(doc, "File: overall.wav   |   Duration: 9.98 s   |   Sample rate: 16 kHz")
    add_bullet(doc, "Language: Mixed multilingual — Telugu, English, and Hindi within a single clip")
    add_bullet(doc, "Used for all Sequential, Concurrent, Poisson, and Stress tests")
    add_bullet(doc, "Multi-duration batch used clips of 2 s – 30 s tiled from the same audio")

    doc.add_paragraph()
    add_heading(doc, "Decoding Strategies Tested", level=2)
    add_bullet(doc, "CTC (Connectionist Temporal Classification)")
    add_bullet(doc, "TDT (Token-and-Duration Transducer)")
    add_bullet(doc, "RNNT (Recurrent Neural Network Transducer)")

    doc.add_paragraph()
    add_heading(doc, "Test Types", level=2)
    add_bullet(doc, "Sequential — 3 warmup + 15 runs, one request at a time")
    add_bullet(doc, "Concurrent — workers = 2, 4, 8 firing simultaneously")
    add_bullet(doc, "Poisson — random arrivals at 4 req/s over a 12 s window")
    add_bullet(doc, "Stress — 200 requests from 100 concurrent workers")
    add_bullet(doc, "Multi-duration Batch — 7 clip lengths (2 s – 30 s) fired concurrently")

    doc.add_paragraph()
    add_heading(doc, "Scenarios Covered", level=2)
    add_body(doc, (
        "17–21 scenarios per strategy: baseline, language constraints (hi, hi_en, te, te_en), "
        "blank penalty, word boost (default/strong/full), keyword boost, timestamps, "
        "timestamps+language, timestamps+word-boost, full pipeline, and masked_joint (mj_) "
        "variants for transducer strategies."
    ))
    doc.add_page_break()

    # ── SECTION 2: DECODING STRATEGIES ──────────────
    add_heading(doc, "2. Decoding Strategies Explained", level=1)
    add_divider(doc)

    add_heading(doc, "CTC — Connectionist Temporal Classification", level=2)
    add_body(doc, (
        "CTC operates by assigning token probabilities to every encoder frame independently. "
        "There is no autoregressive prediction network; all frames are scored in parallel. "
        "This makes CTC extremely fast and highly parallelisable. "
        "Language constraints are applied as a direct -inf vocabulary mask before decoding, "
        "adding only a small overhead from the masking operation. "
        "CTC is the reference baseline in all test runs."
    ))

    doc.add_paragraph()
    add_heading(doc, "TDT — Token-and-Duration Transducer", level=2)
    add_body(doc, (
        "TDT is an autoregressive transducer that jointly predicts both the next token and "
        "the number of frames to advance (duration). By predicting frame skips, TDT avoids "
        "processing every encoder step, making it significantly faster than standard RNNT "
        "while retaining good accuracy. Language constraints can be applied via joint bias "
        "masking (masked_joint / mj_ variants) or post-hoc CTC rescoring. "
        "Run 1 tested CTC + TDT."
    ))

    doc.add_paragraph()
    add_heading(doc, "RNNT — Recurrent Neural Network Transducer", level=2)
    add_body(doc, (
        "RNNT is the standard autoregressive transducer. The prediction network is conditioned "
        "on all previously emitted tokens, making it the most contextually aware and typically "
        "the most accurate decoder — but also the slowest. Language constraints follow the "
        "same mechanism as TDT (masked_joint or legacy joint masking). "
        "Run 2 tested CTC + RNNT."
    ))
    doc.add_page_break()

    # ── SECTION 3: TEST MODES ────────────────────────
    add_heading(doc, "3. Test Modes Explained", level=1)
    add_divider(doc)

    modes = [
        ("Sequential",
         "One request at a time. Three warmup requests followed by 15 timed runs. "
         "Measures clean per-request latency with no queuing effects. The primary metric "
         "for comparing decoding strategy overhead."),
        ("Concurrent",
         "N workers each fire one request simultaneously (N = 2, 4, 8). "
         "Measures how well the server batches parallel requests and how throughput "
         "(req/s) scales with load."),
        ("Poisson",
         "Requests arrive according to a Poisson process at 4 req/s over a 12-second "
         "window (~48 requests). Models realistic, bursty real-world traffic patterns."),
        ("Stress",
         "200 requests launched from 100 concurrent workers simultaneously. "
         "Tests server stability, queue handling, and failure rate under extreme sustained load. "
         "A perfect score is 200/200 successes."),
        ("Multi-duration Batch",
         "Seven audio clips of 2 s, 5 s, 10 s, 15 s, 20 s, 25 s, and 30 s fired both "
         "solo and concurrently. Tests how the dynamic batcher handles mixed-length inputs. "
         "Clips longer than 20 s trigger the chunked transcription path."),
    ]
    for title, body in modes:
        add_heading(doc, title, level=2)
        add_body(doc, body)
        doc.add_paragraph()

    doc.add_page_break()

    # ── SECTION 4: SEQUENTIAL ────────────────────────
    add_heading(doc, "4. Sequential Latency Results", level=1)
    add_divider(doc)
    add_body(doc, (
        "Each test uses 3 warmup requests followed by 15 timed runs on overall.wav (9.98 s). "
        "Latencies are reported in milliseconds. RTF = Real-Time Factor (latency / audio duration). "
        "All RTF values are far below 1.0, confirming real-time capability."
    ))
    doc.add_paragraph()

    # Chart 1
    add_heading(doc, "Chart 1: Sequential Mean Latency — All 17 Common Scenarios", level=2)
    doc.add_picture(chart_paths["chart1"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading(doc, "Run 1 — CTC Sequential", level=2)
    add_body(doc, "Baseline CTC mean latency: 57.2 ms. Full pipeline (lang + word boost + keyword boost + timestamps): 75.8 ms.")
    doc.add_paragraph()
    add_sequential_table(doc, SCENARIOS_COMMON, CTC_RUN1, "CTC Run 1")
    doc.add_paragraph()

    add_heading(doc, "Run 1 — TDT Sequential (includes masked_joint variants)", level=2)
    add_body(doc, "Baseline TDT mean latency: 124.3 ms (~2.2× CTC). Full pipeline: 150.0 ms. masked_joint single-language adds 30–58 ms over legacy.")
    doc.add_paragraph()
    add_sequential_table(doc, TDT_RUN1_SCENARIOS, TDT_RUN1, "TDT Run 1")
    doc.add_paragraph()
    doc.add_page_break()

    add_heading(doc, "Run 2 — CTC Sequential", level=2)
    add_body(doc, "CTC results are highly consistent between runs (56–81 ms), confirming reproducibility.")
    doc.add_paragraph()
    add_sequential_table(doc, SCENARIOS_COMMON, CTC_RUN2, "CTC Run 2")
    doc.add_paragraph()

    add_heading(doc, "Run 2 — RNNT Sequential (includes masked_joint variants)", level=2)
    add_body(doc, "Baseline RNNT mean latency: 179.4 ms (~3.1× CTC, ~1.4× TDT). Full pipeline: 209.9 ms. lang_te is an outlier at 247 ms due to the Telugu vocabulary size and autoregressive overhead.")
    doc.add_paragraph()
    add_sequential_table(doc, RNNT_RUN2_SCENARIOS, RNNT_RUN2, "RNNT Run 2")
    doc.add_page_break()

    add_heading(doc, "Run 3 — CTC Sequential (Optimized)", level=2)
    add_body(doc, "CTC baseline: 60.5 ms. Full pipeline: 66.9 ms. Consistent with previous runs.")
    doc.add_paragraph()
    add_sequential_table(doc, SCENARIOS_COMMON, CTC_RUN3, "CTC Run 3")
    doc.add_paragraph()

    add_heading(doc, "Run 3 — TDT Sequential (Optimized, includes masked_joint variants)", level=2)
    add_body(doc, "Baseline TDT mean latency: 71.8 ms (~1.2x CTC). Full pipeline: 89.0 ms. Optimizations reduced TDT latency by 42% vs Run 1.")
    doc.add_paragraph()
    add_sequential_table(doc, TDT_RUN3_SCENARIOS, TDT_RUN3, "TDT Run 3")
    doc.add_page_break()

    # ── SECTION 5: CONCURRENT ────────────────────────
    add_heading(doc, "5. Concurrent Latency Results", level=1)
    add_divider(doc)
    add_body(doc, (
        "Workers fire requests simultaneously. Mean latency and throughput (req/s) are measured "
        "at concurrency levels 2, 4, and 8."
    ))
    doc.add_paragraph()

    # Chart 2
    add_heading(doc, "Chart 2: Concurrent Latency & Throughput vs Concurrency Level", level=2)
    doc.add_picture(chart_paths["chart2"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Concurrent table
    conc_headers = ["Strategy", "Concurrency 2 — Latency", "Concurrency 2 — Req/s",
                    "Concurrency 4 — Latency", "Concurrency 4 — Req/s",
                    "Concurrency 8 — Latency", "Concurrency 8 — Req/s"]
    ct = doc.add_table(rows=1 + len(CONC_DATA), cols=7)
    ct.style = "Table Grid"
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = ct.rows[0]
    for i, h in enumerate(conc_headers):
        hrow.cells[i].text = h
        set_cell_bg(hrow.cells[i], HEADER_ROW_COLOR)
        set_cell_font(hrow.cells[i], bold=True, color=RGBColor(255, 255, 255), size=8)
        hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, (label, data) in enumerate(CONC_DATA.items()):
        row = ct.rows[r_idx + 1]
        vals = [label,
                f"{data['latency'][0]:.1f} ms", f"{data['throughput'][0]:.2f}",
                f"{data['latency'][1]:.1f} ms", f"{data['throughput'][1]:.2f}",
                f"{data['latency'][2]:.1f} ms", f"{data['throughput'][2]:.2f}"]
        for c_idx, v in enumerate(vals):
            row.cells[c_idx].text = v
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], ALT_ROW_COLOR)
            set_cell_font(row.cells[c_idx], size=9)

    doc.add_page_break()

    # ── SECTION 6: POISSON ───────────────────────────
    add_heading(doc, "6. Poisson Traffic Results", level=1)
    add_divider(doc)
    add_body(doc, (
        "Requests arrive at a Poisson rate over a 12-second window. "
        "Run 1 and Run 2 used 4 req/s (~48 total requests). "
        "Run 3 (Optimized) was tested at 20 req/s (~240 total requests) to stress the optimized pipeline. "
        "Five representative scenarios are reported per strategy."
    ))
    doc.add_paragraph()

    # Chart 3
    add_heading(doc, "Chart 3: Poisson Mean Latency — 5 Key Scenarios", level=2)
    doc.add_picture(chart_paths["chart3"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Poisson table
    pois_headers = ["Strategy", "Baseline", "lang_hi_en", "word_boost", "kw_boost", "full_pipeline"]
    pt = doc.add_table(rows=1 + len(POISSON_DATA), cols=6)
    pt.style = "Table Grid"
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = pt.rows[0]
    for i, h in enumerate(pois_headers):
        hrow.cells[i].text = h
        set_cell_bg(hrow.cells[i], HEADER_ROW_COLOR)
        set_cell_font(hrow.cells[i], bold=True, color=RGBColor(255, 255, 255), size=9)
        hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, (label, vals) in enumerate(POISSON_DATA.items()):
        row = pt.rows[r_idx + 1]
        cells = [label] + [f"{v:.1f} ms" for v in vals]
        for c_idx, v in enumerate(cells):
            row.cells[c_idx].text = v
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], ALT_ROW_COLOR)
            set_cell_font(row.cells[c_idx], size=9)

    doc.add_paragraph()
    add_heading(doc, "Poisson Traffic — Run 3 (Optimized, 20 req/s)", level=2)
    add_body(doc, (
        "Run 3 was tested at 20 req/s (5x higher than Run 1/2) over a 12-second window (~240 total requests)."
    ))
    doc.add_paragraph()

    pt3 = doc.add_table(rows=1 + len(POISSON_DATA_RUN3), cols=6)
    pt3.style = "Table Grid"
    pt3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow3 = pt3.rows[0]
    for i, h in enumerate(pois_headers):
        hrow3.cells[i].text = h
        set_cell_bg(hrow3.cells[i], HEADER_ROW_COLOR)
        set_cell_font(hrow3.cells[i], bold=True, color=RGBColor(255, 255, 255), size=9)
        hrow3.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, (label, vals) in enumerate(POISSON_DATA_RUN3.items()):
        row = pt3.rows[r_idx + 1]
        cells = [label] + [f"{v:.1f} ms" for v in vals]
        for c_idx, v in enumerate(cells):
            row.cells[c_idx].text = v
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], ALT_ROW_COLOR)
            set_cell_font(row.cells[c_idx], size=9)

    doc.add_page_break()

    # ── SECTION 7: STRESS ────────────────────────────
    add_heading(doc, "7. Stress Test Results", level=1)
    add_divider(doc)
    add_body(doc, (
        "200 requests were launched from 100 concurrent workers simultaneously for each strategy. "
        "All strategies achieved 200/200 successful responses with zero failures, "
        "demonstrating server robustness under extreme load."
    ))
    doc.add_paragraph()

    # Chart 4
    add_heading(doc, "Chart 4: Stress Test Latency & Throughput", level=2)
    doc.add_picture(chart_paths["chart4"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Stress table
    st_headers = ["Strategy", "Successes", "Wall Time (s)", "Throughput (req/s)", "Mean Latency (ms)", "P95 Latency (ms)"]
    stt = doc.add_table(rows=1 + len(STRESS_DATA), cols=6)
    stt.style = "Table Grid"
    stt.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = stt.rows[0]
    for i, h in enumerate(st_headers):
        hrow.cells[i].text = h
        set_cell_bg(hrow.cells[i], HEADER_ROW_COLOR)
        set_cell_font(hrow.cells[i], bold=True, color=RGBColor(255, 255, 255), size=9)
        hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, (label, data) in enumerate(STRESS_DATA.items()):
        row = stt.rows[r_idx + 1]
        vals = [label, "200/200", f"{data['wall']:.2f}", f"{data['rate']:.1f}",
                f"{data['mean']}", f"{data['p95']}"]
        for c_idx, v in enumerate(vals):
            row.cells[c_idx].text = v
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], ALT_ROW_COLOR)
            set_cell_font(row.cells[c_idx], size=9)

    doc.add_page_break()

    # ── SECTION 8: LONG-AUDIO RESILIENCE ─────────────
    add_heading(doc, "8. Long-Audio Resilience", level=1)
    add_divider(doc)
    add_body(doc, (
        "Two long-audio clips (60 s and ~599 s / 10 min) were transcribed to verify "
        "the chunked inference path and server stability. All tests completed successfully."
    ))
    doc.add_paragraph()

    la_headers = ["Strategy", "60 s Audio (ms)", "599 s Audio (ms)", "Status"]
    la_data = [
        ("CTC (Run 1)",  "225",   "1 917",  "OK"),
        ("TDT (Run 1)",  "778",   "7 328",  "OK"),
        ("CTC (Run 2)",  "328",   "2 146",  "OK"),
        ("RNNT (Run 2)", "1 135", "10 928", "OK"),
        ("CTC (Run 3)",  "208",   "1 847",  "OK"),
        ("TDT (Run 3)",  "459",   "4 204",  "OK"),
    ]
    lat = doc.add_table(rows=1 + len(la_data), cols=4)
    lat.style = "Table Grid"
    lat.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = lat.rows[0]
    for i, h in enumerate(la_headers):
        hrow.cells[i].text = h
        set_cell_bg(hrow.cells[i], HEADER_ROW_COLOR)
        set_cell_font(hrow.cells[i], bold=True, color=RGBColor(255, 255, 255), size=9)
        hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_vals in enumerate(la_data):
        row = lat.rows[r_idx + 1]
        for c_idx, v in enumerate(row_vals):
            row.cells[c_idx].text = v
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], ALT_ROW_COLOR)
            set_cell_font(row.cells[c_idx], size=9)

    doc.add_paragraph()
    add_body(doc, (
        "CTC handles 10-minute audio in ~1.9 s. TDT takes 7.3 s and RNNT 10.9 s, "
        "both using the chunked transcription path for audio exceeding 20 s. "
        "The linear scaling with audio duration confirms correct chunked processing."
    ))
    doc.add_page_break()

    # ── SECTION 9: MULTI-DURATION ─────────────────────
    add_heading(doc, "9. Multi-Duration Batch Results", level=1)
    add_divider(doc)
    add_body(doc, (
        "Seven audio clips (2 s – 30 s) were processed both in isolation (solo) "
        "and concurrently (all 7 fired simultaneously). Clips exceeding 20 s use the "
        "chunked transcription path."
    ))
    doc.add_paragraph()

    # Chart 5
    add_heading(doc, "Chart 5: Multi-Duration Solo & Concurrent Latency vs Audio Duration", level=2)
    doc.add_picture(chart_paths["chart5"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    def add_multidur_table(doc, data_dict, title, chunked_marker="*"):
        add_heading(doc, title, level=2)
        dur_headers = ["Strategy"] + [f"{d}s" for d in DURATIONS]
        chunked_idx = {25, 30}
        t = doc.add_table(rows=1 + len(data_dict), cols=len(dur_headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hrow = t.rows[0]
        for i, h in enumerate(dur_headers):
            hrow.cells[i].text = h
            set_cell_bg(hrow.cells[i], HEADER_ROW_COLOR)
            set_cell_font(hrow.cells[i], bold=True, color=RGBColor(255, 255, 255), size=8.5)
            hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r_idx, (label, vals) in enumerate(data_dict.items()):
            row = t.rows[r_idx + 1]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(row.cells[0], size=9)
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[0], ALT_ROW_COLOR)
            for c_idx, v in enumerate(vals):
                cell = row.cells[c_idx + 1]
                dur = DURATIONS[c_idx]
                txt = f"{v:.1f}"
                if dur in chunked_idx and title.lower().count("solo") == 0:
                    pass  # concurrent chunking noted elsewhere
                if dur in chunked_idx:
                    txt += " *"
                cell.text = txt
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if r_idx % 2 == 0:
                    set_cell_bg(cell, ALT_ROW_COLOR)
                set_cell_font(cell, size=9)

    add_multidur_table(doc, MULTI_SOLO, "Solo Latency (ms) — One clip at a time")
    doc.add_paragraph()
    add_body(doc, "* = Clips processed via chunked transcription path (audio > 20 s)")
    doc.add_paragraph()
    add_multidur_table(doc, MULTI_CONC, "Concurrent Latency (ms) — All 7 clips fired simultaneously")
    doc.add_page_break()

    # ── SECTION 10: MASKED_JOINT ──────────────────────
    add_heading(doc, "10. masked_joint vs Legacy Language Constraint", level=1)
    add_divider(doc)
    add_body(doc, (
        "The masked_joint (mj_) variants apply language constraints directly in the joint "
        "network via per-item serial decoding. This is the correct approach for transducers "
        "but incurs additional latency compared to the legacy approach."
    ))
    doc.add_paragraph()

    # Chart 6
    add_heading(doc, "Chart 6: masked_joint vs Legacy — TDT & RNNT", level=2)
    doc.add_picture(chart_paths["chart6"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    mj_headers = ["Constraint", "TDT Legacy (ms)", "TDT masked_joint (ms)", "TDT Delta",
                  "RNNT Legacy (ms)", "RNNT masked_joint (ms)", "RNNT Delta"]
    mj_rows = [
        ("lang_hi",    143.1, 182.5, "+39.4", 196.4, 202.0, "+5.6"),
        ("lang_hi_en", 141.4, 129.7, "-11.7", 198.9, 184.2, "-14.7"),
        ("lang_te",    149.9, 180.5, "+30.6", 247.1, 253.7, "+6.6"),
        ("lang_te_en", 132.6, 129.7,  "-2.9", 186.9, 182.8, "-4.1"),
    ]
    mjt = doc.add_table(rows=1 + len(mj_rows), cols=7)
    mjt.style = "Table Grid"
    mjt.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = mjt.rows[0]
    for i, h in enumerate(mj_headers):
        hrow.cells[i].text = h
        set_cell_bg(hrow.cells[i], HEADER_ROW_COLOR)
        set_cell_font(hrow.cells[i], bold=True, color=RGBColor(255, 255, 255), size=8.5)
        hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_vals in enumerate(mj_rows):
        row = mjt.rows[r_idx + 1]
        for c_idx, v in enumerate(row_vals):
            row.cells[c_idx].text = str(v)
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], ALT_ROW_COLOR)
            set_cell_font(row.cells[c_idx], size=9)

    doc.add_paragraph()
    add_body(doc, (
        "Key finding: masked_joint adds 30–40 ms for single-language constraints (hi, te) because "
        "serial per-item decoding is forced. For multi-language constraints (hi_en, te_en) "
        "the overhead is negligible or even slightly negative (speedup from tighter beam), "
        "because most vocabulary tokens remain valid."
    ))
    doc.add_page_break()

    # ── SECTION 11: OPTIMIZATIONS APPLIED ────────────
    add_heading(doc, "11. Optimizations Applied (Run 1 \u2192 Run 3)", level=1)
    add_divider(doc)
    add_body(doc, (
        "The following optimizations were applied between Run 1 (unoptimized TDT) and "
        "Run 3 (optimized TDT) to reduce per-request latency and improve throughput:"
    ))
    doc.add_paragraph()

    optimizations = [
        ("Batched masked_joint decode",
         "Language-constrained items in RNNT/TDT are now grouped by (lang_tags, blank_penalty) "
         "and decoded in a single batched call instead of per-item B=1 serial decoding. "
         "This eliminates the per-item decode overhead that caused 30-50ms per item."),
        ("Pre-decode CTC word_boost overlap",
         "CTC word-level language boosting (ctc_decode_with_lang_boost) is submitted to a "
         "thread pool during the pre-decode callback, so CPU-bound word_boost runs in parallel "
         "with GPU TDT/RNNT decode (~10ms/item hidden behind GPU time)."),
        ("Pre-decode keyword spotting overlap",
         "run_word_spotter() is also submitted to the thread pool during the pre-decode callback. "
         "The expensive CPU beam search (~20ms/item) runs during GPU decode time. After decode, "
         "only the fast merge step runs."),
        ("Word spotter hot-loop optimizations",
         "Token class changed from @dataclass to __slots__ (saves ~200 bytes and ~30% creation "
         "time per instance, millions created per request). beam_pruning changed from np.argmax "
         "on Python list to pure-Python max() generator (avoids numpy array allocation per frame). "
         "make_boost_fn factory returns mode-specialized closures, eliminating 40+ string "
         "comparisons per call in the inner loop."),
        ("Triton dynamic batcher tuning",
         "max_batch_size reduced from 32 to 16, queue_delay set to 20ms, preferred_batch_size "
         "[2,4,8,16]. RNNT sub-batch size increased from 4 to 8."),
    ]

    for i, (title, detail) in enumerate(optimizations, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}")
        run.bold = True
        run.font.color.rgb = HEADING_COLOR
        run.font.size = Pt(11)
        add_body(doc, detail)
        doc.add_paragraph()

    # Before vs After comparison table
    add_heading(doc, "Before vs After Comparison", level=2)
    doc.add_paragraph()

    ba_headers = ["Metric", "TDT (Run 1)", "TDT (Run 3)", "Improvement"]
    ba_rows = [
        ("Sequential baseline", "124.3 ms", "71.8 ms", "42% faster"),
        ("Sequential full_pipeline", "150.0 ms", "89.0 ms", "41% faster"),
        ("mj_lang_hi (masked joint)", "182.5 ms", "92.2 ms", "49% faster"),
        ("mj_lang_te", "180.5 ms", "96.2 ms", "47% faster"),
        ("Concurrent 8 latency", "344.5 ms", "166.3 ms", "52% faster"),
        ("Concurrent 8 throughput", "23.1 req/s", "41.8 req/s", "81% higher"),
        ("Poisson baseline*", "147.7 ms @4rps", "103.6 ms @20rps", "5x higher load, 30% lower latency"),
        ("Poisson full_pipeline*", "187.6 ms @4rps", "206.1 ms @20rps", "5x higher load, similar latency"),
        ("Stress throughput", "22.2 req/s", "48.7 req/s", "120% higher"),
        ("60s audio", "778 ms", "459 ms", "41% faster"),
    ]

    bat = doc.add_table(rows=1 + len(ba_rows), cols=4)
    bat.style = "Table Grid"
    bat.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = bat.rows[0]
    for i, h in enumerate(ba_headers):
        hrow.cells[i].text = h
        set_cell_bg(hrow.cells[i], HEADER_ROW_COLOR)
        set_cell_font(hrow.cells[i], bold=True, color=RGBColor(255, 255, 255), size=9)
        hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_vals in enumerate(ba_rows):
        row = bat.rows[r_idx + 1]
        for c_idx, v in enumerate(row_vals):
            row.cells[c_idx].text = v
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], ALT_ROW_COLOR)
            set_cell_font(row.cells[c_idx], size=9)

    doc.add_paragraph()
    add_body(doc, "*Note: Run 3 Poisson was tested at 20 req/s (5x higher than Run 1's 4 req/s)")
    doc.add_paragraph()

    # Chart 7
    add_heading(doc, "Chart 7: TDT Optimization Impact — Run 1 vs Run 3", level=2)
    doc.add_picture(chart_paths["chart7"], width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── SECTION 12: ANALYSIS ──────────────────────────
    add_heading(doc, "12. Analysis & Key Insights", level=1)
    add_divider(doc)

    insights = [
        ("Optimized TDT is now only ~1.2× CTC (down from 2.2×), closing the latency gap",
         "Before optimization: CTC baseline 57 ms, TDT baseline 124 ms (2.2× gap). "
         "After optimization: CTC baseline 61 ms, TDT baseline 72 ms (1.2× gap). "
         "The remaining 11 ms gap is the irreducible cost of autoregressive TDT decoding. "
         "RNNT baseline remains at 179 ms (3.1× CTC) as it was not optimized in Run 3."),

        ("TDT full_pipeline latency reduced by 41% (150 ms → 89 ms)",
         "The full pipeline (lang_tags + word_boost + keyword_boost + timestamps + masked_joint) "
         "dropped from 150 ms to 89 ms. This was achieved by overlapping CPU-bound post-processing "
         "(word_boost, keyword spotting) with GPU TDT decode, and batching masked_joint items "
         "instead of per-item B=1 serial decoding."),

        ("masked_joint overhead eliminated: mj_lang_hi dropped from 183 ms to 92 ms (49% faster)",
         "Batched masked_joint decode groups items by (lang_tags, blank_penalty) for a single "
         "batched decode call. Before: per-item B=1 serial decoding added 30-50 ms per item. "
         "After: batched decode adds only 3-5 ms over baseline. Single-language Telugu constraint "
         "(mj_lang_te) dropped from 181 ms to 96 ms (47% faster)."),

        ("TDT now handles 20 req/s Poisson traffic (5× the original 4 req/s test rate)",
         "Run 3 Poisson was tested at 20 req/s. TDT baseline: 104 ms mean, 150 ms P95. "
         "TDT full_pipeline at 20 req/s: 206 ms mean, 341 ms P95. "
         "For comparison, Run 1 TDT at only 4 req/s showed 148 ms baseline and 188 ms full_pipeline. "
         "The optimized server handles 5× the load with comparable or better latencies."),

        ("TDT stress throughput doubled: 22 → 49 req/s",
         "Under stress (200 requests, 100 concurrent workers), TDT throughput increased from "
         "22.2 to 48.7 req/s — nearly matching CTC's 48.6 req/s. Mean stress latency dropped "
         "from 3695 ms to 1582 ms. This demonstrates that the decode-time overlap and batched "
         "masked_joint optimizations compound under high concurrency."),

        ("All strategies remain far below real-time (RTF << 1.0)",
         "CTC RTF = 0.006 (170× real-time), optimized TDT RTF = 0.007–0.011, RNNT RTF = 0.018–0.025. "
         "Even the most expensive optimized TDT scenario (lang_te at 106 ms) completes a 10 s clip "
         "in under 110 ms."),

        ("Language constraining cost reduced: 5–35 ms on optimized TDT (was 15–70 ms)",
         "CTC vocab masking is cheap (5–10 ms). Optimized TDT with batched masked_joint: "
         "multi-language (hi_en, te_en) adds only 5–6 ms; single-language (hi, te) adds 20–34 ms. "
         "Telugu single-language remains the most expensive due to larger effective vocabulary, "
         "but is now 96 ms vs the original 181 ms."),

        ("Keyword boosting adds ~16–18 ms, overlapped with GPU decode on TDT path",
         "The word-spotter beam search runs on CTC log-probabilities. On the TDT path, "
         "keyword spotting is now submitted to a thread pool during the pre-decode callback, "
         "hiding the ~20 ms CPU cost behind ~30-40 ms GPU decode time. The residual serial "
         "cost is only the fast merge step (~1-2 ms)."),

        ("CTC throughput scales to 51 req/s at concurrency 8; optimized TDT reaches 42 req/s",
         "CTC concurrent throughput: 39 → 47 → 51 req/s at concurrency 2/4/8. "
         "Optimized TDT: 31 → 30 → 42 req/s. The TDT dip at concurrency 4 is due to "
         "sub-batch boundaries (MAX_RNNT_SUB_BATCH=8). At concurrency 8 the full sub-batch "
         "is utilized, recovering throughput."),

        ("Stress test: 200/200 requests succeeded for all strategies — zero failures",
         "Even under 100 concurrent workers, zero failures were observed across CTC, TDT, and RNNT. "
         "Optimized TDT stress mean latency (1582 ms) is now comparable to CTC (1532 ms), "
         "down from 3695 ms before optimization."),

        ("Long audio: optimized TDT handles 10 min in 4.2 s (was 7.3 s, 43% faster)",
         "All strategies correctly activate the chunked transcription path for audio > 20 s. "
         "CTC: 60s→208ms, 599s→1847ms. Optimized TDT: 60s→459ms, 599s→4204ms. "
         "RNNT: 60s→1135ms, 599s→10928ms. Latency scales linearly with duration."),
    ]

    for i, (headline, detail) in enumerate(insights, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {headline}")
        run.bold = True
        run.font.color.rgb = HEADING_COLOR
        run.font.size = Pt(11)
        add_body(doc, detail)
        doc.add_paragraph()

    doc.add_page_break()

    # ── SECTION 13: SUMMARY TABLE ─────────────────────
    add_heading(doc, "13. Summary Comparison", level=1)
    add_divider(doc)
    add_body(doc, "High-level comparison of all decoding strategies across key dimensions, including the optimized TDT (Run 3).")
    doc.add_paragraph()

    sum_headers = ["Metric", "CTC", "TDT", "RNNT", "TDT (Optimized)"]
    sum_rows = [
        ("Sequential baseline (ms)", "57", "124", "179", "72"),
        ("Sequential full_pipeline (ms)", "76\u201381", "150", "210", "89"),
        ("RTF (baseline)", "0.006", "0.012", "0.018", "0.007"),
        ("Relative speed (vs CTC)", "1.0\u00d7", "~2.2\u00d7", "~3.1\u00d7", "~1.2\u00d7"),
        ("Concurrent 8 \u2014 latency (ms)", "219\u2013243", "345", "454", "166"),
        ("Concurrent 8 \u2014 throughput (req/s)", "32\u201336", "23", "18", "42"),
        ("Poisson baseline (ms)", "59\u201362", "148", "220", "104 @20rps"),
        ("Stress mean latency (ms)", "2 031\u20132 115", "3 695", "4 667", "1 582"),
        ("Stress throughput (req/s)", "38\u201340", "22", "18", "48.7"),
        ("60 s audio (ms)", "225\u2013328", "778", "1 135", "459"),
        ("599 s audio (ms)", "1 917\u20132 146", "7 328", "10 928", "4 204"),
        ("Failure rate (stress)", "0%", "0%", "0%", "0%"),
        ("Language constraint cost", "5\u201317 ms", "15\u201370 ms", "15\u201370 ms", "5\u201335 ms"),
        ("Keyword boost cost", "~12\u201317 ms", "~12\u201317 ms", "~12\u201317 ms", "~16\u201318 ms"),
    ]

    smt = doc.add_table(rows=1 + len(sum_rows), cols=5)
    smt.style = "Table Grid"
    smt.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = smt.rows[0]
    for i, h in enumerate(sum_headers):
        hrow.cells[i].text = h
        set_cell_bg(hrow.cells[i], HEADER_ROW_COLOR)
        set_cell_font(hrow.cells[i], bold=True, color=RGBColor(255, 255, 255), size=10)
        hrow.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row_vals in enumerate(sum_rows):
        row = smt.rows[r_idx + 1]
        for c_idx, v in enumerate(row_vals):
            row.cells[c_idx].text = v
            row.cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 0:
                set_cell_bg(row.cells[c_idx], ALT_ROW_COLOR)
            set_cell_font(row.cells[c_idx], size=10)

    doc.add_paragraph()
    add_heading(doc, "Recommendation", level=2)
    add_body(doc, (
        "CTC remains the recommended strategy for latency-critical production deployments. "
        "It is the fastest decoder, scales best under concurrency, "
        "and handles keyword/language constraints with minimal overhead. "
        "However, with the Run 3 optimizations, TDT is now only ~1.2x slower than CTC "
        "(down from 2.2x in Run 1), making it a viable option for latency-critical production "
        "when higher transcription accuracy is needed. Optimized TDT achieves 42 req/s at "
        "concurrency 8 (vs 23 req/s unoptimized) and handles 5x higher Poisson load at "
        "comparable latencies. "
        "RNNT is best suited for offline or batch processing where maximum accuracy "
        "is the priority and latency is less constrained."
    ))

    return doc


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    print("Generating charts...")
    chart_paths = {
        "chart1": chart1_sequential_bar(),
        "chart2": chart2_concurrent(),
        "chart3": chart3_poisson(),
        "chart4": chart4_stress(),
        "chart5": chart5_multiduration(),
        "chart6": chart6_maskedjoint(),
        "chart7": chart7_optimization_comparison(),
    }
    for name, path in chart_paths.items():
        print(f"  {name}: {path}")

    print("Building document...")
    doc = build_doc(chart_paths)

    out_path = "/home/nobroker-tlt415/Documents/claude/combined_implementation/latency_report.docx"
    doc.save(out_path)
    size = os.path.getsize(out_path)
    print(f"\nSaved: {out_path}  ({size:,} bytes)")
    return out_path


if __name__ == "__main__":
    main()

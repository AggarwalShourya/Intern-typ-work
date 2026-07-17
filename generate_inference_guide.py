"""
generate_inference_guide.py
Generates inference_guide.docx — a client-side inference technical guide
for the nemo_hybrid Triton gRPC ASR server.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ────────────────────────────────────────────────────────────
BLUE_HEADING   = RGBColor(0x1F, 0x49, 0x7D)   # dark navy-blue
BLUE_ACCENT    = RGBColor(0x2E, 0x74, 0xB5)   # mid-blue (h2)
BLUE_ACCENT3   = RGBColor(0x1E, 0x5E, 0x9E)   # h3 colour
ROW_LIGHT      = RGBColor(0xDE, 0xEB, 0xF7)   # alternating table row (light blue)
ROW_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
HDR_BG         = RGBColor(0x1F, 0x49, 0x7D)   # table header background
HDR_FG         = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG        = RGBColor(0xF2, 0xF2, 0xF2)   # code-block background
TIP_BG         = RGBColor(0xEB, 0xF3, 0xFB)   # tip/note box background
WARN_BG        = RGBColor(0xFF, 0xF0, 0xE0)   # warning box background

MONO_FONT = "Courier New"
BODY_FONT = "Calibri"


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    """Set cell shading colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = str(rgb)  # RGBColor.__str__ returns 6-char hex e.g. "1F497D"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=True, bottom=True, left=True, right=True,
                    color="B8CCE4", size="4"):
    """Add thin borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, flag in [("top", top), ("bottom", bottom),
                       ("left", left), ("right", right)]:
        if flag:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_para_bg(para, rgb: RGBColor):
    """Add paragraph-level shading (for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_color = str(rgb)  # RGBColor.__str__ returns 6-char hex
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_para_border(para, color="B8CCE4"):
    """Add a left bar border to a paragraph (for callout boxes)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


# ── Document-level helpers ────────────────────────────────────────────────────

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE_HEADING
    run.font.name = BODY_FONT
    # horizontal rule below
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E74B5")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE_ACCENT
    run.font.name = BODY_FONT
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE_ACCENT3
    run.font.name = BODY_FONT
    return p


def add_body(doc, text, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = BODY_FONT
    run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    if italic:
        run.italic = True
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = BODY_FONT
    run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    return p


def add_inline_code(para, text):
    """Append an inline monospace run to an existing paragraph."""
    run = para.add_run(text)
    run.font.name = MONO_FONT
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    return run


def add_code_block(doc, code_text, size_pt=9):
    """Render a grey code block (one paragraph per line)."""
    lines = code_text.split("\n")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        set_para_bg(p, CODE_BG)
        run = p.add_run(line if line else " ")
        run.font.name = MONO_FONT
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # small spacer after block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(6)


def add_note_box(doc, text, bg=TIP_BG, label="Note", border_color="2E74B5"):
    """A shaded info/tip/warning box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    set_para_bg(p, bg)
    set_para_border(p, color=border_color)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.name = BODY_FONT
    label_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    body_run = p.add_run(text)
    body_run.font.size = Pt(10)
    body_run.font.name = BODY_FONT
    body_run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)


def add_styled_table(doc, headers, rows, col_widths=None):
    """
    Create a styled table with a blue header and alternating body rows.
    headers: list of str
    rows:    list of lists of str
    col_widths: list of Inches values (optional)
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, HDR_BG)
        set_cell_border(cell, color="1F497D")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = BODY_FONT
        run.font.color.rgb = HDR_FG

    # Body rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = ROW_LIGHT if r_idx % 2 == 0 else ROW_WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="B8CCE4")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Pt(3)
            # Detect inline code spans wrapped in backticks
            parts = cell_text.split("`")
            for j, part in enumerate(parts):
                if j % 2 == 1:
                    run = p.add_run(part)
                    run.font.name = MONO_FONT
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
                else:
                    # Handle newlines inside non-code parts
                    subparts = part.split("\n")
                    for k, subpart in enumerate(subparts):
                        if k > 0:
                            run = p.add_run("\n")
                            run.font.size = Pt(9.5)
                            run.font.name = BODY_FONT
                        run = p.add_run(subpart)
                        run.font.size = Pt(9.5)
                        run.font.name = BODY_FONT
                        run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_page_break(doc):
    doc.add_page_break()


# =============================================================================
# CONTENT
# =============================================================================

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title page ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(40)
    title_para.paragraph_format.space_after  = Pt(6)
    tr = title_para.add_run("nemo_hybrid ASR Server")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = BLUE_HEADING
    tr.font.name = BODY_FONT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after  = Pt(4)
    sr = subtitle.add_run("Client Inference Guide")
    sr.bold = True
    sr.font.size = Pt(20)
    sr.font.color.rgb = BLUE_ACCENT
    sr.font.name = BODY_FONT

    # Decorative rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after  = Pt(8)
    rr = rule.add_run("\u2500" * 60)
    rr.font.color.rgb = BLUE_ACCENT
    rr.font.size = Pt(10)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "Triton gRPC  \u00b7  nemo_hybrid model  \u00b7  Multilingual Indian-language ASR\n"
        "Audience: Python developers with the Triton client installed"
    )
    mr.font.size = Pt(11)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    mr.font.name = BODY_FONT

    add_page_break(doc)

    # =========================================================================
    # 1. OVERVIEW
    # =========================================================================
    add_heading1(doc, "1  Overview")

    add_body(doc,
        "The nemo_hybrid Triton Inference Server exposes a high-performance, "
        "multilingual Automatic Speech Recognition (ASR) service over gRPC. "
        "The server hosts a NeMo hybrid CTC/RNNT/TDT model and accepts raw 16 kHz "
        "int16 PCM audio, returning UTF-8 transcripts, log-probability confidence "
        "scores, and optional word-level timestamps.")

    add_body(doc,
        "All feature toggles (language constraining, keyword boosting, word-level "
        "boosting, timestamps, decoding strategy) are controlled entirely from the "
        "client via standard Triton inference inputs -- no server restart or "
        "configuration change is needed.")

    add_heading2(doc, "Supported Languages")
    add_body(doc,
        "The model supports the following BCP-47-style language tags for "
        "vocabulary constraining and per-word language boosting:")

    add_styled_table(doc,
        headers=["Tag", "Language"],
        rows=[
            ["`en`", "English"],
            ["`hi`", "Hindi"],
            ["`ta`", "Tamil"],
            ["`bn`", "Bengali"],
            ["`te`", "Telugu"],
            ["`kn`", "Kannada"],
            ["`mr`", "Marathi"],
            ["`ml`", "Malayalam"],
            ["`gu`", "Gujarati"],
        ],
        col_widths=[Inches(1.2), Inches(2.5)])

    add_heading2(doc, "Decoding Strategies")
    add_body(doc,
        "The server supports four decoding strategies, switchable at runtime via "
        "set_decoding_strategy(). All strategies support the same feature set "
        "(language constraining, keyword boosting, word-level boosting, timestamps) "
        "unless explicitly noted in the Parameter Reference.")

    add_styled_table(doc,
        headers=["Strategy", "Description"],
        rows=[
            ["`\"ctc\"`",
             "CTC greedy batch decode. Fastest. Fully parallel, no autoregressive step. "
             "blank_penalty applies only to this path."],
            ["`\"rnnt\"`",
             "RNNT greedy batch decode. Autoregressive transducer."],
            ["`\"malsd_batch\"`",
             "RNNT MALSD beam decode. Default out-of-the-box. Best accuracy for RNNT models."],
            ["`\"tdt\"`",
             "TDT beam decode. Use for TDT-trained model checkpoints."],
        ],
        col_widths=[Inches(1.8), Inches(5.0)])

    add_note_box(doc,
        "The server auto-corrects incompatible strategy requests. For example, "
        "requesting \"rnnt\" on a TDT-trained model will silently switch to \"tdt\" "
        "and log a warning. You do not need to catch this case explicitly.",
        label="Info")

    add_page_break(doc)

    # =========================================================================
    # 2. SETUP
    # =========================================================================
    add_heading1(doc, "2  Setup")

    add_body(doc,
        "The following imports, constants, and client object are assumed throughout "
        "this guide. Create the client once at module level and reuse it for all "
        "requests.")

    add_code_block(doc, """\
import json
import numpy as np
import soundfile as sf
import tritonclient.grpc as grpcclient
from typing import Any, Dict, List, Optional

# -- Configuration ------------------------------------------------------------
SERVER_URL  = "localhost:9020"   # Triton gRPC address
MODEL_NAME  = "nemo_hybrid"
GRAPH_MODEL = "context_graph_builder"
SAMPLE_RATE = 16000

# -- Shared client (create once, reuse) ---------------------------------------
client = grpcclient.InferenceServerClient(url=SERVER_URL, verbose=False)""")

    add_body(doc,
        "Change SERVER_URL to point at the host running the Triton server "
        "(e.g. \"192.168.1.10:9020\"). The client is thread-safe and can be shared "
        "across concurrent inference calls.")

    add_page_break(doc)

    # =========================================================================
    # 3. HELPER FUNCTIONS
    # =========================================================================
    add_heading1(doc, "3  Helper Functions")

    # 3.1 load_audio_int16
    add_heading2(doc, "3.1  load_audio_int16()")
    add_body(doc,
        "Loads an audio file from disk, converts it to mono 16 kHz float32, then "
        "scales and clips to int16. This is the correct format required by the "
        "AUDIO input tensor.")

    add_code_block(doc, """\
def load_audio_int16(path: str, target_sr: int = SAMPLE_RATE) -> np.ndarray:
    \"\"\"Load audio file, resample if needed, return int16 PCM numpy array.\"\"\"
    audio, sr = sf.read(path, dtype="float32", always_2d=False)
    if audio.ndim > 1:
        audio = audio.mean(axis=1)          # stereo -> mono
    if sr != target_sr:
        import librosa
        audio = librosa.resample(audio, orig_sr=sr, target_sr=target_sr)
    return (audio * 32767).clip(-32768, 32767).astype(np.int16)""")

    add_note_box(doc,
        "librosa is imported lazily -- it is only required if the source file's "
        "sample rate differs from 16000 Hz. WAV files recorded at 16 kHz can be "
        "loaded with soundfile alone.",
        label="Note")

    # 3.2 build_context_graph
    add_heading2(doc, "3.2  build_context_graph()")
    add_body(doc,
        "Sends a list of keyword strings to the context_graph_builder model on the "
        "same Triton server. The server tokenizes the keywords using the model's own "
        "SentencePiece tokenizer, builds a ContextGraphCTC prefix tree, serializes "
        "it to bytes, and returns it. The returned bytes can be passed directly to "
        "infer() as the context_graph argument. No local tokenizer or NeMo model "
        "installation is required -- all tokenization happens server-side.")

    add_code_block(doc, """\
def build_context_graph(keywords: list) -> bytes:
    \"\"\"
    Ask the server to tokenize keywords and build a ContextGraphCTC.
    Returns pickled bytes for use as context_graph in infer().
    No local tokenizer or NeMo model required.
    \"\"\"
    kw_json = json.dumps(keywords).encode()
    inp = grpcclient.InferInput("KEYWORDS", [1], "BYTES")
    inp.set_data_from_numpy(np.array([kw_json], dtype=object))
    out = grpcclient.InferRequestedOutput("CONTEXT_GRAPH")
    result = client.infer(GRAPH_MODEL, [inp], outputs=[out])
    return bytes(result.as_numpy("CONTEXT_GRAPH")[0])""")

    add_note_box(doc,
        "Build the context graph once per unique keyword list and cache the bytes "
        "object. Passing the same bytes to multiple infer() calls is safe and "
        "avoids unnecessary round-trips to the server.",
        label="Performance tip")

    # 3.3 set_decoding_strategy
    add_heading2(doc, "3.3  set_decoding_strategy()")
    add_body(doc,
        "Switches the server's active decoding strategy. This is implemented by "
        "sending a short silent audio clip alongside the SET_DECODING_STRATEGY "
        "control input. The silent clip is required because it must exceed the "
        "mel-filterbank convolution padding (> 512 samples); 8000 samples (0.5 s) "
        "is always safe.")

    add_code_block(doc, """\
def set_decoding_strategy(strategy: str) -> None:
    \"\"\"
    Switch the server's decoding strategy at runtime.
    Valid values: "ctc", "rnnt", "malsd_batch", "tdt"
    \"\"\"
    _N = 8000  # 0.5 s -- must exceed mel-filterbank conv padding (>512 samples)
    dummy = np.zeros((1, _N), dtype=np.int16)
    inp_audio = grpcclient.InferInput("AUDIO", [1, _N], "INT16")
    inp_audio.set_data_from_numpy(dummy)

    inp_strat = grpcclient.InferInput("SET_DECODING_STRATEGY", [1, 1], "BYTES")
    inp_strat.set_data_from_numpy(np.array([[strategy.encode()]], dtype=object))

    outputs = [
        grpcclient.InferRequestedOutput("TRANSCRIPT"),
        grpcclient.InferRequestedOutput("SCORE"),
        grpcclient.InferRequestedOutput("WORD_TIMESTAMPS"),
    ]
    client.infer(MODEL_NAME, [inp_audio, inp_strat], outputs=outputs)
    print(f"[client] Decoding strategy -> {strategy}")""")

    add_note_box(doc,
        "WARNING: set_decoding_strategy() changes the server's global state. "
        "The new strategy applies to ALL subsequent requests from ALL clients until "
        "changed again. Do not use this function concurrently or in production "
        "environments where multiple callers share the same server instance. "
        "If you need per-request strategy control, run separate server instances.",
        label="WARNING",
        bg=WARN_BG,
        border_color="E67E22")

    add_page_break(doc)

    # =========================================================================
    # 4. THE infer() FUNCTION
    # =========================================================================
    add_heading1(doc, "4  The infer() Function")

    add_body(doc,
        "infer() is the single entry-point for all transcription requests. All "
        "feature flags are passed as keyword arguments -- omit any argument to "
        "use its default (disabled/off).")

    add_code_block(doc, """\
def infer(
    audio:             np.ndarray,
    lang_tags:         Optional[List[str]] = None,
    blank_penalty:     float               = 0.0,
    cb_params:         Optional[Dict]      = None,
    context_graph:     Optional[bytes]     = None,
    word_boost_params: Optional[Dict]      = None,
    timestamps:        bool                = False,
    use_masked_joint:  bool                = False,
    timeout_s:         Optional[float]     = None,
) -> Dict[str, Any]:
    inputs = []
    outputs = [
        grpcclient.InferRequestedOutput("TRANSCRIPT"),
        grpcclient.InferRequestedOutput("SCORE"),
        grpcclient.InferRequestedOutput("WORD_TIMESTAMPS"),
    ]

    audio_np = audio.reshape(1, -1)
    inp = grpcclient.InferInput("AUDIO", list(audio_np.shape), "INT16")
    inp.set_data_from_numpy(audio_np)
    inputs.append(inp)

    if lang_tags is not None:
        val = np.array([[json.dumps(lang_tags).encode()]], dtype=object)
        inp = grpcclient.InferInput("LANG_TAGS", [1, 1], "BYTES")
        inp.set_data_from_numpy(val)
        inputs.append(inp)

    if blank_penalty != 0.0:
        val = np.array([[blank_penalty]], dtype=np.float32)
        inp = grpcclient.InferInput("BLANK_PENALTY", [1, 1], "FP32")
        inp.set_data_from_numpy(val)
        inputs.append(inp)

    if cb_params is not None:
        val = np.array([[json.dumps(cb_params).encode()]], dtype=object)
        inp = grpcclient.InferInput("CB_PARAMS", [1, 1], "BYTES")
        inp.set_data_from_numpy(val)
        inputs.append(inp)

    if context_graph is not None:
        val = np.array([[context_graph]], dtype=object)
        inp = grpcclient.InferInput("CONTEXT_GRAPH", [1, 1], "BYTES")
        inp.set_data_from_numpy(val)
        inputs.append(inp)

    if word_boost_params is not None:
        val = np.array([[json.dumps(word_boost_params).encode()]], dtype=object)
        inp = grpcclient.InferInput("WORD_BOOST_PARAMS", [1, 1], "BYTES")
        inp.set_data_from_numpy(val)
        inputs.append(inp)

    if timestamps:
        val = np.array([[True]], dtype=bool)
        inp = grpcclient.InferInput("TIMESTAMPS", [1, 1], "BOOL")
        inp.set_data_from_numpy(val)
        inputs.append(inp)

    if use_masked_joint:
        val = np.array([[True]], dtype=bool)
        inp = grpcclient.InferInput("USE_MASKED_JOINT", [1, 1], "BOOL")
        inp.set_data_from_numpy(val)
        inputs.append(inp)

    result = client.infer(MODEL_NAME, inputs, outputs=outputs,
                          client_timeout=timeout_s)

    transcript = result.as_numpy("TRANSCRIPT")[0, 0]
    if isinstance(transcript, bytes):
        transcript = transcript.decode("utf-8")
    score = float(result.as_numpy("SCORE")[0, 0])

    wt_raw = result.as_numpy("WORD_TIMESTAMPS")[0, 0]
    if isinstance(wt_raw, bytes):
        wt_raw = wt_raw.decode("utf-8")
    word_timestamps = json.loads(wt_raw) if wt_raw else []

    return {"transcript": transcript, "score": score,
            "word_timestamps": word_timestamps}""")

    # Quick-reference table
    add_heading2(doc, "Parameter Quick-Reference")
    add_styled_table(doc,
        headers=["Parameter", "Type", "Default", "Purpose"],
        rows=[
            ["`audio`", "`np.ndarray` (int16)", "Required", "16 kHz mono PCM audio"],
            ["`lang_tags`", "`List[str]` or `None`", "`None`",
             "BCP-47 tags to restrict vocabulary"],
            ["`blank_penalty`", "`float`", "`0.0`",
             "CTC blank logit subtraction (CTC path only)"],
            ["`cb_params`", "`Dict` or `None`", "`None`",
             "Keyword-boosting hyperparameters (requires `context_graph`)"],
            ["`context_graph`", "`bytes` or `None`", "`None`",
             "Serialized ContextGraphCTC (requires `cb_params`)"],
            ["`word_boost_params`", "`Dict` or `None`", "`None`",
             "Per-word language boosting on CTC log-probs"],
            ["`timestamps`", "`bool`", "`False`",
             "Enable word-level timestamp output"],
            ["`use_masked_joint`", "`bool`", "`False`",
             "Mask joint layer during RNNT/TDT decode (with `lang_tags`)"],
            ["`timeout_s`", "`float` or `None`", "`None`",
             "Client-side RPC timeout in seconds"],
        ],
        col_widths=[Inches(1.6), Inches(1.6), Inches(1.0), Inches(3.5)])

    add_page_break(doc)

    # =========================================================================
    # 5. PARAMETER REFERENCE
    # =========================================================================
    add_heading1(doc, "5  Parameter Reference")

    add_body(doc,
        "This section documents every infer() parameter in detail. Each entry "
        "covers: what the parameter does, the accepted values, when and why to "
        "use it, interactions with other parameters, and practical tuning tips.")

    # -- 5.1 audio ------------------------------------------------------------
    add_heading2(doc, "5.1  audio  (required)")

    add_body(doc,
        "A 1-D NumPy array of int16 PCM samples recorded at 16 kHz, mono. "
        "This is the only required argument. The function automatically reshapes "
        "it to shape (1, N) before sending to the server.")

    p = add_body(doc, "Triton tensor:  ")
    add_inline_code(p, "AUDIO [1, N]  INT16")

    add_body(doc, "Preparation:", space_after=2)
    add_code_block(doc, """\
# From a file (recommended):
audio_int16 = load_audio_int16("speech.wav")

# From a float32 array (e.g. from a streaming buffer):
audio_int16 = (float_audio * 32767).clip(-32768, 32767).astype(np.int16)""")

    add_note_box(doc,
        "The audio must be at least ~512 samples (32 ms at 16 kHz) to satisfy "
        "the mel-filterbank convolution padding requirement inside the encoder. "
        "Shorter clips will cause a server-side error. The set_decoding_strategy() "
        "helper uses 8000 samples (0.5 s) as a safe minimum for control calls.",
        label="Constraint")

    # -- 5.2 lang_tags --------------------------------------------------------
    add_heading2(doc, "5.2  lang_tags  (Optional[List[str]], default None)")

    add_body(doc,
        "Hard-restricts the decoded vocabulary to tokens belonging to the listed "
        "language(s). Accepted values are BCP-47-style tags: "
        "\"en\", \"hi\", \"ta\", \"bn\", \"te\", \"kn\", \"mr\", \"ml\", \"gu\".")

    add_body(doc, "When to use:", italic=True, space_after=2)
    add_bullet(doc,
        "You know the language(s) present in the audio in advance.")
    add_bullet(doc,
        "The audio contains code-switching (e.g., Hindi-English) and you want to "
        "prevent spurious tokens from unrelated scripts.")
    add_bullet(doc,
        "You want to force a specific-language output regardless of acoustic "
        "ambiguity.")

    add_body(doc, "Behaviour per decoding strategy:", italic=True, space_after=2)

    add_styled_table(doc,
        headers=["Strategy mode", "Behaviour"],
        rows=[
            ["CTC (`\"ctc\"`)",
             "Applies a hard -inf mask on the logit vector for all tokens not belonging "
             "to the listed languages, before greedy argmax decode. Fully deterministic."],
            ["RNNT / TDT  (default, `use_masked_joint=False`)",
             "Decodes unconstrained first (full speed, fully batched). Then, any token "
             "not belonging to the listed languages is replaced post-hoc using "
             "the CTC log-probability sequence at the corresponding emission frame. "
             "Fast, but may fail for very short or acoustically ambiguous utterances."],
            ["RNNT / TDT  (`use_masked_joint=True`)",
             "Applies a -inf bias mask to the joint output layer before each decode "
             "step, ensuring the prediction network is always conditioned on valid "
             "in-language tokens. Most accurate for short/ambiguous utterances, "
             "at the cost of per-item serial decode (increased latency)."],
        ],
        col_widths=[Inches(2.2), Inches(5.5)])

    add_body(doc, "Examples:", italic=True, space_after=2)
    add_code_block(doc, """\
# Single language
result = infer(audio_int16, lang_tags=["hi"])

# Bilingual (Hindi + English code-switching)
result = infer(audio_int16, lang_tags=["hi", "en"])

# Omit to allow all supported languages (default)
result = infer(audio_int16)""")

    add_note_box(doc,
        "LANG_TAGS and WORD_BOOST_PARAMS are independent and can be combined. "
        "LANG_TAGS applies a hard constraint (removes tokens entirely) while "
        "WORD_BOOST_PARAMS applies a soft boost. Using both together is valid and "
        "often produces the best results for multilingual audio.",
        label="Interaction")

    # -- 5.3 blank_penalty ----------------------------------------------------
    add_heading2(doc, "5.3  blank_penalty  (float, default 0.0)")

    add_body(doc,
        "Subtracts this value from the CTC blank logit before softmax, which "
        "forces the CTC decoder to emit more non-blank tokens. This parameter "
        "applies only to the CTC decoding path -- it is ignored for RNNT and TDT "
        "strategies.")

    add_body(doc, "When to use:", italic=True, space_after=2)
    add_bullet(doc,
        "Combined with lang_tags to prevent the model from staying in the blank "
        "state and emitting a near-empty transcript.")
    add_bullet(doc,
        "When language-constrained decoding produces very short or missing output "
        "despite clear speech -- a blank_penalty of 2.0-4.0 typically fixes this.")
    add_bullet(doc,
        "When you need more aggressive token emission for short utterances.")

    add_body(doc, "Accepted values:", italic=True, space_after=2)
    add_styled_table(doc,
        headers=["Value", "Effect"],
        rows=[
            ["`0.0`", "Disabled (default). No modification to blank logit."],
            ["`1.0 - 2.0`", "Light penalty. Useful for longer utterances with lang_tags."],
            ["`2.0 - 4.0`",
             "Moderate. Typical range when combining with language constraining. "
             "Start here when lang_tags produces empty output."],
            ["`4.0 - 5.0`",
             "Aggressive. Forces many token emissions. Use only for very short "
             "utterances or if lower values still produce empty/truncated output."],
            ["`> 5.0`",
             "Rarely needed. May cause hallucination artifacts on silence or noise."],
        ],
        col_widths=[Inches(1.4), Inches(6.3)])

    add_code_block(doc, """\
# Typical usage with language constraining:
result = infer(audio_int16, lang_tags=["hi"], blank_penalty=3.0)
result = infer(audio_int16, lang_tags=["hi", "en"], blank_penalty=3.0)""")

    add_note_box(doc,
        "blank_penalty is ignored on RNNT and TDT strategies. However, "
        "word_boost_params has its own separate blank_penalty key that operates "
        "on the CTC log-prob sequence post-encoder -- that one does apply "
        "regardless of decoding strategy.",
        label="CTC only")

    # -- 5.4 cb_params + context_graph ----------------------------------------
    add_heading2(doc, "5.4  cb_params  (Optional[Dict], default None)")
    add_heading3(doc, "5.5  context_graph  (Optional[bytes], default None)")

    add_body(doc,
        "These two parameters must be provided together to activate keyword "
        "boosting. context_graph is the serialized ContextGraphCTC prefix tree "
        "(built once via build_context_graph()) and cb_params is a JSON dictionary "
        "of boosting hyperparameters.")

    add_body(doc,
        "Pass cb_params={} to use all server-side defaults. Any subset of keys "
        "can be overridden in the dict.")

    add_body(doc, "cb_params keys:", italic=True, space_after=2)
    add_styled_table(doc,
        headers=["Key", "Type", "Default", "Description"],
        rows=[
            ["`beam_threshold`", "`float`", "`17.0`",
             "Prune beams whose score falls more than this below the best beam. "
             "Lower = more aggressive pruning (faster, may lose keywords). "
             "Raise if keywords are being dropped."],
            ["`cb_weight`", "`float`", "`2.5`",
             "Primary keyword boost strength. Higher = stronger bias toward keyword "
             "tokens. Typical range 1.5-4.0."],
            ["`keyword_threshold`", "`float`", "`5.0`",
             "Minimum accumulated CTC log-prob score for a keyword path to be "
             "considered active. Lower = more sensitive detection."],
            ["`blank_threshold`", "`float`", "`0.13`",
             "CTC blank probability above which the frame is considered a blank "
             "frame (no token emission expected)."],
            ["`non_blank_threshold`", "`float`", "`0.002`",
             "Minimum non-blank probability to consider a token as a candidate at "
             "a given frame."],
            ["`ctc_ali_token_weight`", "`float`", "`0.5`",
             "Weight applied to CTC alignment token scores when combining with the "
             "beam graph score."],
            ["`intersection_threshold`", "`float`", "`30.0`",
             "Controls overlap detection between competing keyword hypotheses. "
             "Lower = stricter overlap resolution."],
            ["`cb_weight_continuation`", "`float`", "`2.7`",
             "Boost weight for the continuation of an already-accepted keyword "
             "sub-path. Slightly higher than cb_weight to stabilize long keywords."],
            ["`boost_mode`", "`str`", "`\"entropy_weighted\"`",
             "How the boost is applied. entropy_weighted scales boost by frame-level "
             "entropy (less boost on confident frames, more on ambiguous frames)."],
            ["`keyword_overlap_threshold`", "`float`", "`30.0`",
             "Score threshold for resolving overlapping keyword detections when "
             "resolve_keyword_overlaps is true."],
            ["`resolve_keyword_overlaps`", "`bool`", "`true`",
             "When true, overlapping keyword detections are resolved to keep the "
             "highest-scoring non-overlapping set."],
        ],
        col_widths=[Inches(2.0), Inches(0.7), Inches(1.0), Inches(4.0)])

    add_body(doc, "Examples:", italic=True, space_after=2)
    add_code_block(doc, """\
# Build the context graph once, reuse across requests:
context_graph_bytes = build_context_graph(["keyword", "boosting", "target phrase"])

# Use all defaults:
result = infer(audio_int16, cb_params={}, context_graph=context_graph_bytes)

# Increase boost strength and lower beam threshold:
result = infer(
    audio_int16,
    cb_params={"cb_weight": 3.5, "beam_threshold": 12.0, "keyword_threshold": 4.0},
    context_graph=context_graph_bytes,
)

# Combine with language constraining:
result = infer(
    audio_int16,
    lang_tags=["te", "en"],
    cb_params={"cb_weight": 2.5},
    context_graph=context_graph_bytes,
)""")

    add_note_box(doc,
        "cb_params and context_graph must always be supplied together. Providing "
        "one without the other has no effect. Build the context graph once per "
        "unique keyword list and cache the bytes object -- it is safe to reuse "
        "across concurrent calls.",
        label="Important")

    # -- 5.6 word_boost_params ------------------------------------------------
    add_heading2(doc, "5.6  word_boost_params  (Optional[Dict], default None)")

    add_body(doc,
        "Applies soft per-word language boosting on the CTC log-probability frame "
        "sequence after the encoder forward pass but before final decoding. For "
        "each decoded word, the server identifies its dominant language from the "
        "CTC probs, then adds lang_boost to that language's token log-probs and "
        "subtracts lang_penalty from all other language tokens (excluding blank).")

    add_body(doc,
        "This mechanism is independent of LANG_TAGS and operates on a different "
        "level: LANG_TAGS enforces a hard constraint on the vocabulary; "
        "word_boost_params applies a soft, per-word directional push.")

    add_body(doc, "word_boost_params keys:", italic=True, space_after=2)
    add_styled_table(doc,
        headers=["Key", "Type", "Default", "Description"],
        rows=[
            ["`enabled`", "`bool`", "`false`",
             "Must be true to activate. If false or omitted, no boosting occurs."],
            ["`lang_boost`", "`float`", "`2.0`",
             "Additive boost (in log-prob space) applied to tokens belonging to the "
             "dominant language of each word. Higher = stronger language preference. "
             "Typical range: 1.5-3.0."],
            ["`lang_penalty`", "`float`", "`0.0`",
             "Subtractive penalty applied to tokens belonging to non-dominant languages "
             "(blank excluded). Can be used alone or together with lang_boost. "
             "Typical range: 0.0-2.0. Setting both creates a push-pull effect."],
            ["`blank_penalty`", "`float`", "`0.5`",
             "Penalty subtracted from the CTC blank log-prob (post-encoder). Applied "
             "regardless of decoding strategy -- unlike the top-level blank_penalty "
             "parameter which is CTC-only."],
        ],
        col_widths=[Inches(1.7), Inches(0.7), Inches(1.0), Inches(4.3)])

    add_body(doc, "Examples:", italic=True, space_after=2)
    add_code_block(doc, """\
# Minimal: activate with all defaults
result = infer(audio_int16, word_boost_params={"enabled": True})

# Boost only (push dominant language up, leave others unchanged)
result = infer(audio_int16, word_boost_params={
    "enabled": True, "lang_boost": 2.5, "blank_penalty": 0.4
})

# Boost + penalty (push-pull effect for stronger language separation)
result = infer(audio_int16, word_boost_params={
    "enabled": True, "lang_boost": 2.0, "lang_penalty": 1.5, "blank_penalty": 0.4
})

# Combined with lang_tags for maximum control:
result = infer(
    audio_int16,
    lang_tags=["te"],
    word_boost_params={"enabled": True, "lang_boost": 2.0,
                       "lang_penalty": 1.5, "blank_penalty": 0.5},
)""")

    add_note_box(doc,
        "lang_boost and lang_penalty are complementary and can be used "
        "independently or together. Using both creates a push-pull effect: "
        "lang_boost raises the dominant language's tokens while lang_penalty lowers "
        "competing languages. Start with lang_boost alone (e.g. 2.0) and add "
        "lang_penalty (e.g. 1.0-1.5) only if you need stronger separation.",
        label="Tuning tip")

    # -- 5.7 timestamps -------------------------------------------------------
    add_heading2(doc, "5.7  timestamps  (bool, default False)")

    add_body(doc,
        "When True, the server computes word-level start and end timestamps and "
        "returns them in the word_timestamps field of the response. When False "
        "(default), word_timestamps is an empty list.")

    add_body(doc, "How timestamps are computed:", italic=True, space_after=2)
    add_styled_table(doc,
        headers=["Strategy", "Method"],
        rows=[
            ["CTC (`\"ctc\"`)",
             "Greedy argmax on the CTC log-prob sequence, followed by CTC collapse "
             "(merge repeated tokens, remove blanks). The start/end of each word's "
             "token span in the frame sequence is converted to seconds using the "
             "model's time-stride."],
            ["RNNT / TDT",
             "Per-token encoder-frame indices are recorded from the hypothesis object "
             "during beam/greedy decode. Tokens are grouped into words by the "
             "SentencePiece word-boundary marker (\u25be). Start/end of each word's "
             "frame range is converted to seconds."],
        ],
        col_widths=[Inches(2.0), Inches(5.7)])

    add_body(doc, "Output format:", italic=True, space_after=2)
    add_code_block(doc, """\
# Each entry in word_timestamps:
{"word": str, "start_s": float, "end_s": float}

# Example output:
[
    {"word": "hello",  "start_s": 0.040, "end_s": 0.360},
    {"word": "world",  "start_s": 0.400, "end_s": 0.720},
]

# Pretty-print:
for w in result["word_timestamps"]:
    print(f"{w['word']:<25} {w['start_s']:>6.3f}s  {w['end_s']:>6.3f}s")""")

    add_note_box(doc,
        "timestamps=True can be combined with all other parameters (lang_tags, "
        "word_boost_params, cb_params / context_graph, use_masked_joint). There "
        "is no additional latency cost for timestamps on the CTC path. On the "
        "RNNT/TDT path, frame indices are tracked during the existing decode loop.",
        label="Note")

    # -- 5.8 use_masked_joint -------------------------------------------------
    add_heading2(doc, "5.8  use_masked_joint  (bool, default False)")

    add_body(doc,
        "Controls how LANG_TAGS are applied during RNNT and TDT decoding. Has no "
        "effect on the CTC strategy or when lang_tags is None.")

    add_styled_table(doc,
        headers=["Value", "Behaviour"],
        rows=[
            ["`False` (default)",
             "Post-hoc replacement: decode unconstrained (full batch speed), then "
             "replace any out-of-language tokens in the hypothesis using CTC log-probs "
             "at the corresponding encoder frame. Fast and works well for longer "
             "utterances (>= 1 s). May fail for very short or acoustically ambiguous "
             "utterances where the prediction network was conditioned on wrong tokens."],
            ["`True`",
             "Masked-joint: apply a -inf bias mask to the joint output layer at each "
             "decode step, so the prediction network is always conditioned on valid "
             "in-language tokens. The most correct approach for short utterances or "
             "when language is highly ambiguous acoustically. Requires per-item "
             "serial decode for constrained items -- adds latency proportional to the "
             "number of constrained items in the batch."],
        ],
        col_widths=[Inches(1.4), Inches(6.3)])

    add_body(doc, "When to use use_masked_joint=True:", italic=True, space_after=2)
    add_bullet(doc,
        "Short utterances (< 0.5 s) where the default post-hoc approach produces "
        "wrong-language output.")
    add_bullet(doc,
        "Languages that are acoustically ambiguous at the beginning of the "
        "utterance (e.g., Telugu vs. Kannada).")
    add_bullet(doc,
        "When accuracy is more important than latency and the batch is small.")

    add_code_block(doc, """\
# Switch to TDT (or RNNT) strategy first:
set_decoding_strategy("tdt")

# Default post-hoc (fast, good for longer audio):
result = infer(audio_int16, lang_tags=["te"], blank_penalty=3.0)

# Masked-joint (best for short/ambiguous):
result = infer(audio_int16, lang_tags=["te"], blank_penalty=3.0, use_masked_joint=True)

# Mixed language with masked-joint:
result = infer(audio_int16, lang_tags=["te", "en"], blank_penalty=2.0,
               use_masked_joint=True)""")

    # -- 5.9 timeout_s --------------------------------------------------------
    add_heading2(doc, "5.9  timeout_s  (Optional[float], default None)")

    add_body(doc,
        "Client-side gRPC timeout in seconds. If the server does not respond "
        "within this duration, the Triton client raises an InferenceServerException "
        "with a DEADLINE_EXCEEDED status. Set to None (default) for no timeout.")

    add_code_block(doc, """\
# Fail fast if inference takes longer than 5 seconds:
try:
    result = infer(audio_int16, timeout_s=5.0)
except Exception as e:
    print(f"Inference timed out: {e}")""")

    add_note_box(doc,
        "timeout_s is a client-side deadline only -- the server will continue "
        "processing after the client disconnects. For server-side request "
        "cancellation, use Triton's built-in request cancellation API separately.",
        label="Note")

    add_page_break(doc)

    # =========================================================================
    # 6. RETURN VALUE
    # =========================================================================
    add_heading1(doc, "6  Return Value")

    add_body(doc,
        "infer() always returns a Python dict with three keys, regardless of which "
        "optional parameters were supplied:")

    add_code_block(doc, """\
{
    "transcript":      str,          # UTF-8 transcription string
    "score":           float,        # log-probability confidence score (<= 0.0)
    "word_timestamps": List[Dict],   # [] when timestamps=False
}""")

    add_styled_table(doc,
        headers=["Key", "Type", "Description"],
        rows=[
            ["`\"transcript\"`", "`str`",
             "The recognised UTF-8 text. Empty string if speech was not detected. "
             "Supports all Unicode scripts (Devanagari, Tamil, Telugu, etc.)."],
            ["`\"score\"`", "`float`",
             "Log-probability confidence score (<= 0.0). Higher (less negative) means "
             "higher model confidence. Not normalised by length -- longer utterances "
             "tend to have lower (more negative) scores. Use for relative ranking, "
             "not absolute thresholding."],
            ["`\"word_timestamps\"`", "`List[Dict]`",
             "List of {\"word\": str, \"start_s\": float, \"end_s\": float} entries. "
             "Empty list when timestamps=False. Times are in seconds from the start "
             "of the audio array."],
        ],
        col_widths=[Inches(1.8), Inches(1.0), Inches(4.9)])

    add_page_break(doc)

    # =========================================================================
    # 7. ASYNC INFERENCE
    # =========================================================================
    add_heading1(doc, "7  Async Inference")

    add_body(doc,
        "The synchronous infer() function blocks the calling thread until the "
        "server responds. For non-blocking inference, two async variants are "
        "provided \u2014 infer_async() (native asyncio, awaitable) and infer_future() "
        "(callback-based Future, works in synchronous/threaded code). Both accept "
        "exactly the same parameters as infer() and return the same "
        "{transcript, score, word_timestamps} dict.")

    # -- 7.1 Setup ------------------------------------------------------------
    add_heading2(doc, "7.1  Setup")

    add_body(doc, "Both clients must be created at startup:")

    add_code_block(doc, """\
import asyncio
import concurrent.futures
import tritonclient.grpc     as grpcclient
import tritonclient.grpc.aio as grpcclient_aio

client     = grpcclient.InferenceServerClient(url=SERVER_URL, verbose=False)
aio_client = grpcclient_aio.InferenceServerClient(url=SERVER_URL)""")

    # -- 7.2 Choosing the Right Variant ---------------------------------------
    add_heading2(doc, "7.2  Choosing the Right Variant")

    add_styled_table(doc,
        headers=["Variant", "Transport", "When to use"],
        rows=[
            ["`infer_async()`",
             "`tritonclient.grpc.aio` (native asyncio)",
             "Code is already `async` \u2014 FastAPI endpoints, aiohttp handlers, asyncio "
             "scripts. Each `await` yields control so other coroutines can run "
             "concurrently."],
            ["`infer_future()`",
             "`tritonclient.grpc` callback + `concurrent.futures.Future`",
             "Synchronous or threaded code. Fires the gRPC request immediately; "
             "calling thread is free to do other work before calling `.result()`. "
             "Also bridges into asyncio via `asyncio.wrap_future()`."],
        ],
        col_widths=[Inches(1.5), Inches(2.5), Inches(3.7)])

    # -- 7.3 infer_async() ---------------------------------------------------
    add_heading2(doc, "7.3  infer_async() \u2014 native asyncio")

    add_code_block(doc, """\
async def infer_async(
    audio:             np.ndarray,
    lang_tags:         Optional[List[str]] = None,
    blank_penalty:     float               = 0.0,
    cb_params:         Optional[Dict]      = None,
    context_graph:     Optional[bytes]     = None,
    word_boost_params: Optional[Dict]      = None,
    timestamps:        bool                = False,
    use_masked_joint:  bool                = False,
    timeout_s:         Optional[float]     = None,
) -> Dict[str, Any]:
    inputs, outputs = _build_inputs_outputs(
        audio, lang_tags, blank_penalty, cb_params,
        context_graph, word_boost_params, timestamps, use_masked_joint,
    )
    result = await aio_client.infer(
        MODEL_NAME, inputs, outputs=outputs, client_timeout=timeout_s
    )
    return _parse_result(result)""")

    add_body(doc, "Usage examples:", italic=True, space_after=2)
    add_code_block(doc, """\
# Single request
result = await infer_async(audio_int16, lang_tags=["hi"])

# Fan out multiple requests concurrently -- all inflight at the same time
results = await asyncio.gather(
    infer_async(clip1, lang_tags=["hi"]),
    infer_async(clip2, lang_tags=["te"]),
    infer_async(clip3),
)

# Inside a FastAPI endpoint
@app.post("/transcribe")
async def transcribe(audio_bytes: bytes):
    audio = np.frombuffer(audio_bytes, dtype=np.int16)
    result = await infer_async(audio, lang_tags=["hi", "en"])
    return result""")

    # -- 7.4 infer_future() --------------------------------------------------
    add_heading2(doc, "7.4  infer_future() \u2014 concurrent.futures.Future")

    add_code_block(doc, """\
def infer_future(
    audio:             np.ndarray,
    lang_tags:         Optional[List[str]] = None,
    blank_penalty:     float               = 0.0,
    cb_params:         Optional[Dict]      = None,
    context_graph:     Optional[bytes]     = None,
    word_boost_params: Optional[Dict]      = None,
    timestamps:        bool                = False,
    use_masked_joint:  bool                = False,
    timeout_s:         Optional[float]     = None,
) -> concurrent.futures.Future:
    inputs, outputs = _build_inputs_outputs(
        audio, lang_tags, blank_penalty, cb_params,
        context_graph, word_boost_params, timestamps, use_masked_joint,
    )
    future = concurrent.futures.Future()

    def _callback(result, error):
        if error is not None:
            future.set_exception(error)
        else:
            try:
                future.set_result(_parse_result(result))
            except Exception as exc:
                future.set_exception(exc)

    client.async_infer(
        MODEL_NAME, inputs, callback=_callback,
        outputs=outputs, client_timeout=timeout_s,
    )
    return future""")

    add_body(doc, "Usage examples:", italic=True, space_after=2)
    add_code_block(doc, """\
# Fire and collect later -- calling thread not blocked in between
fut = infer_future(audio_int16, lang_tags=["hi"])
# ... do other work here ...
result = fut.result()   # blocks only at this point

# Fan out many requests, collect as they complete
futures = [infer_future(clip, lang_tags=["hi"]) for clip in clips]
for fut in concurrent.futures.as_completed(futures):
    result = fut.result()
    print(result["transcript"])

# Bridge a Future into asyncio when you need to await it
result = await asyncio.wrap_future(infer_future(audio_int16))""")

    add_note_box(doc,
        "infer_future() uses the synchronous client.async_infer() callback "
        "mechanism, which internally uses a background network thread managed by "
        "the gRPC channel. The calling thread returns immediately \u2014 no thread pool "
        "or executor is needed on the client side.",
        label="Note")

    # -- 7.5 Comparison -------------------------------------------------------
    add_heading2(doc, "7.5  Comparison")

    add_styled_table(doc,
        headers=["Property", "infer()", "infer_async()", "infer_future()"],
        rows=[
            ["Blocks calling thread", "Yes", "No (yields to event loop)", "No (returns immediately)"],
            ["Requires event loop",   "No",  "Yes",                       "No"],
            ["Fan-out idiom",         "ThreadPoolExecutor", "asyncio.gather()", "as_completed()"],
            ["Bridge to asyncio",     "\u2014", "Native",                  "asyncio.wrap_future()"],
            ["Timeout parameter",     "timeout_s", "timeout_s",           "timeout_s"],
        ],
        col_widths=[Inches(2.0), Inches(1.2), Inches(1.8), Inches(2.7)])

    add_page_break(doc)

    # =========================================================================
    # 8. FEATURE COMBINATIONS
    # =========================================================================
    add_heading1(doc, "8  Feature Combinations")

    add_body(doc,
        "Most features are orthogonal and can be stacked freely. The table below "
        "summarises which combinations are valid, which are meaningful, and any "
        "constraints.")

    add_styled_table(doc,
        headers=["Combination", "Valid?", "Notes"],
        rows=[
            ["`lang_tags` + `word_boost_params`",
             "Yes",
             "Hard vocab mask (lang_tags) + soft per-word push (word_boost_params). "
             "Combining both gives maximum control: lang_tags removes impossible tokens, "
             "word_boost_params amplifies the correct language within the allowed set."],
            ["`lang_tags` + `blank_penalty`",
             "Yes (CTC only)",
             "blank_penalty forces more token emissions so lang_tags does not produce "
             "an empty transcript. Use blank_penalty=2.0-4.0 when lang_tags causes "
             "missing output on CTC."],
            ["`lang_tags` + `use_masked_joint`",
             "Yes (RNNT/TDT only)",
             "use_masked_joint strengthens lang_tags on RNNT/TDT. Only meaningful when "
             "a non-CTC strategy is active."],
            ["`cb_params` + `context_graph`",
             "Required pair",
             "Must be provided together. Keyword boosting is inactive unless both are "
             "supplied."],
            ["`cb_params` + `lang_tags`",
             "Yes",
             "Keyword boosting + language constraining. The keyword paths are matched "
             "within the language-constrained vocabulary."],
            ["`cb_params` + `word_boost_params`",
             "Yes",
             "Keyword boosting (CTC graph) and word-level language boosting are "
             "independent mechanisms and can both be active simultaneously."],
            ["`timestamps` + any feature",
             "Yes",
             "Timestamps can be combined with all other parameters. No additional "
             "latency on CTC; minimal overhead on RNNT/TDT."],
            ["`blank_penalty` on RNNT/TDT",
             "No effect",
             "The top-level blank_penalty parameter is silently ignored on RNNT and TDT "
             "strategies. Use word_boost_params[\"blank_penalty\"] instead, which "
             "operates post-encoder regardless of strategy."],
            ["`use_masked_joint` without `lang_tags`",
             "No effect",
             "use_masked_joint is ignored when lang_tags is None."],
        ],
        col_widths=[Inches(2.5), Inches(1.0), Inches(4.2)])

    add_page_break(doc)

    # =========================================================================
    # 9. COMPLETE EXAMPLES
    # =========================================================================
    add_heading1(doc, "9  Complete Examples")

    add_body(doc,
        "All examples below assume audio_int16 has already been loaded via "
        "load_audio_int16() and the client is initialised as shown in Section 2.")

    # Test 1
    add_heading2(doc, "Test 1 -- Pure Transcription (Baseline)")
    add_body(doc,
        "The simplest possible call: send audio, get a transcript back. "
        "No language constraining, no boosting, no timestamps. "
        "Uses whatever decoding strategy the server is currently configured with.")
    add_code_block(doc, """\
result = infer(audio_int16)
print(result["transcript"])
print(f"Confidence score: {result['score']:.4f}")""")

    # Test 2a
    add_heading2(doc, "Test 2a -- Language-Constrained: Single Language")
    add_body(doc,
        "Restrict the output vocabulary to Hindi tokens only. The blank_penalty "
        "of 3.0 prevents the CTC decoder from staying in the blank state when "
        "the vocabulary is heavily restricted.")
    add_code_block(doc, """\
result = infer(audio_int16, lang_tags=["hi"], blank_penalty=3.0)
print(result["transcript"])""")

    # Test 2b
    add_heading2(doc, "Test 2b -- Language-Constrained: Bilingual")
    add_body(doc,
        "Allow both Hindi and English tokens. Useful for code-switching audio. "
        "The blank_penalty remains important to prevent blank-dominated output "
        "on the CTC path.")
    add_code_block(doc, """\
result = infer(audio_int16, lang_tags=["hi", "en"], blank_penalty=3.0)
print(result["transcript"])""")

    # Test 2c
    add_heading2(doc, "Test 2c -- Masked-Joint Language Constraining (RNNT/TDT)")
    add_body(doc,
        "For short or acoustically ambiguous utterances, use_masked_joint=True "
        "ensures the RNNT/TDT prediction network is always conditioned on valid "
        "in-language tokens. First switch the server to TDT (or RNNT) strategy, "
        "then compare the two approaches.")
    add_code_block(doc, """\
set_decoding_strategy("tdt")

# Post-hoc approach (fast, works for longer audio):
result_legacy = infer(audio_int16, lang_tags=["te"], blank_penalty=3.0)

# Masked-joint (best for short/ambiguous utterances):
result_mj = infer(audio_int16, lang_tags=["te"], blank_penalty=3.0,
                  use_masked_joint=True)

# Bilingual masked-joint:
result_mj_mixed = infer(audio_int16, lang_tags=["te", "en"], blank_penalty=2.0,
                        use_masked_joint=True)

print("Post-hoc:     ", result_legacy["transcript"])
print("Masked-joint: ", result_mj["transcript"])
print("Mixed MJ:     ", result_mj_mixed["transcript"])""")

    # Test 3
    add_heading2(doc, "Test 3 -- Word-Level Language Boosting")
    add_body(doc,
        "word_boost_params applies a soft per-word language boost on the CTC "
        "log-prob sequence. Unlike lang_tags, this does not remove any tokens -- "
        "it just nudges probabilities toward the dominant language of each word.")
    add_code_block(doc, """\
# 3a -- activate with all defaults:
result = infer(audio_int16, word_boost_params={"enabled": True})

# 3b -- stronger boost only (no cross-language penalty):
result = infer(audio_int16, word_boost_params={
    "enabled": True, "lang_boost": 2.5, "blank_penalty": 0.4
})

# 3c -- boost + penalty (push-pull for stronger language separation):
result = infer(audio_int16, word_boost_params={
    "enabled": True, "lang_boost": 2.0, "lang_penalty": 1.5, "blank_penalty": 0.4
})""")

    # Test 4
    add_heading2(doc, "Test 4 -- LANG_TAGS + Word-Level Boosting Combined")
    add_body(doc,
        "Combine both mechanisms: lang_tags applies a hard vocabulary constraint "
        "while word_boost_params adds a soft directional boost within that "
        "constrained space. This typically gives the cleanest results for "
        "single-language audio.")
    add_code_block(doc, """\
result = infer(
    audio_int16,
    lang_tags=["te"],
    word_boost_params={
        "enabled": True,
        "lang_boost": 2.0,
        "lang_penalty": 1.5,
        "blank_penalty": 0.5,
    },
)
print(result["transcript"])""")

    # Test 5
    add_heading2(doc, "Test 5 -- Keyword Boosting")
    add_body(doc,
        "Build a ContextGraphCTC from a keyword list (server-side tokenization, "
        "no local NeMo required), then pass it alongside cb_params. "
        "The context graph bytes can be reused across many infer() calls.")
    add_code_block(doc, """\
# Build once, reuse:
context_graph_bytes = build_context_graph(["keyword", "boosting"])

# 5a -- use all server defaults:
result = infer(audio_int16, cb_params={}, context_graph=context_graph_bytes)

# 5b -- custom hyperparameters:
result = infer(
    audio_int16,
    cb_params={
        "cb_weight": 3.0,
        "keyword_threshold": 4.0,
        "beam_threshold": 15.0,
    },
    context_graph=context_graph_bytes,
)
print(result["transcript"])""")

    # Test 6
    add_heading2(doc, "Test 6 -- Keyword Boosting + Language Constraining")
    add_body(doc,
        "Combine keyword boosting with language constraining. The keyword paths "
        "are matched within the language-constrained vocabulary, so keywords "
        "from the specified language(s) are boosted while out-of-language tokens "
        "remain suppressed.")
    add_code_block(doc, """\
result = infer(
    audio_int16,
    lang_tags=["te", "en"],
    cb_params={"cb_weight": 2.5},
    context_graph=context_graph_bytes,
)
print(result["transcript"])""")

    # Test 7-9
    add_heading2(doc, "Tests 7-9 -- Word-Level Timestamps")
    add_body(doc,
        "Add timestamps=True to any infer() call to receive word-level start "
        "and end times. Timestamps can be combined with all other features.")
    add_code_block(doc, """\
# Test 7: Pure CTC timestamps (simplest):
result = infer(audio_int16, timestamps=True)

# Test 8: Language constraining + word boost + timestamps:
result = infer(
    audio_int16,
    lang_tags=["hi", "en"],
    word_boost_params={"enabled": True, "lang_boost": 2.0, "blank_penalty": 0.4},
    timestamps=True,
)

# Test 9: Print word-level timestamp table:
print(f"Transcript: {result['transcript']}")
print(f"Score:      {result['score']:.4f}")
print()
print(f"{'Word':<25} {'Start':>8}  {'End':>8}")
print("-" * 46)
for w in result["word_timestamps"]:
    print(f"{w['word']:<25} {w['start_s']:>6.3f}s  {w['end_s']:>6.3f}s")""")

    add_page_break(doc)

    # =========================================================================
    # 10. QUICK REFERENCE CARD
    # =========================================================================
    add_heading1(doc, "10  Quick Reference Card")

    add_body(doc,
        "A compact summary of all features, their associated parameters, and "
        "recommended starting values.")

    add_styled_table(doc,
        headers=["Goal", "Parameters", "Example"],
        rows=[
            ["Basic transcription",
             "audio only",
             "`infer(audio_int16)`"],
            ["Single-language output",
             "`lang_tags=[\"xx\"]`\n`blank_penalty=3.0` (CTC)",
             "`infer(audio, lang_tags=[\"hi\"], blank_penalty=3.0)`"],
            ["Bilingual output",
             "`lang_tags=[\"xx\", \"yy\"]`\n`blank_penalty=3.0` (CTC)",
             "`infer(audio, lang_tags=[\"hi\", \"en\"], blank_penalty=3.0)`"],
            ["Short utterance lang constraint",
             "`lang_tags=[...]`\n`use_masked_joint=True`",
             "Switch to RNNT/TDT first, then:\n"
             "`infer(audio, lang_tags=[\"te\"], use_masked_joint=True)`"],
            ["Soft language boost",
             "`word_boost_params={\"enabled\":True, ...}`",
             "`infer(audio, word_boost_params={\"enabled\":True, \"lang_boost\":2.0})`"],
            ["Keyword boosting",
             "`cb_params={...}`\n`context_graph=<bytes>`",
             "Build once: `build_context_graph([...])`\n"
             "Then: `infer(audio, cb_params={}, context_graph=graph_bytes)`"],
            ["Word timestamps",
             "`timestamps=True`",
             "`infer(audio, timestamps=True)`"],
            ["Switch decoding strategy",
             "`set_decoding_strategy(\"ctc\")`\nor `\"rnnt\"`, `\"malsd_batch\"`, `\"tdt\"`",
             "Call before infer(). Server-global change."],
            ["Request timeout",
             "`timeout_s=5.0`",
             "`infer(audio, timeout_s=5.0)`"],
        ],
        col_widths=[Inches(1.8), Inches(2.2), Inches(3.7)])

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "nemo_hybrid ASR Server \u00b7 Client Inference Guide  \u00b7  "
        "All examples use the gRPC Triton client (tritonclient.grpc)")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    fr.font.name = BODY_FONT
    fr.italic = True

    return doc


if __name__ == "__main__":
    import os
    out_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "inference_guide.docx"
    )
    doc = build_document()
    doc.save(out_path)
    print(f"Document saved to: {out_path}")
